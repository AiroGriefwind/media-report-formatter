# =============================================================================
# SHARED WEB SCRAPING FUNCTIONS
# =============================================================================

import time
import base64
import tempfile
import os
import traceback
from functools import wraps
from collections import defaultdict
import re

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from twocaptcha import TwoCaptcha
from docx import Document
from datetime import datetime

# Import shared Wisers functions
from .wisers_utils import (
    retry_step,
    wait_for_search_results,
    scroll_to_load_all_content,
    wait_for_ajax_complete,
    wait_for_enabled_search_button,
    wait_for_results_panel_ready,
    ensure_results_list_visible,
    inject_cjk_font_css,
    set_media_filters_in_panel,
)
from .html_structure_config import HTML_STRUCTURE

from .config import WISERS_URL, MEDIA_NAME_MAPPINGS, EDITORIAL_MEDIA_ORDER, EDITORIAL_MEDIA_NAMES
from utils.firebase_logging import get_logger

# -----------------------------------------------------------------
# Helper: detect the “no-article” empty-state in multiple layouts
# -----------------------------------------------------------------
TAB_BAR_ZEROS_XPATH = (
    "//ul[contains(@class,'nav-tabs') and contains(@class,'navbar-nav-pub')]"
    "/li//span[normalize-space(text())='(0)']"
)

def _detect_no_article_banner(driver):
    """
    Return True if a 'no article' or 'no data' banner exists.
    """
    try:
        els = driver.find_elements(
            By.XPATH,
            "//h5[contains(text(),'没有文章') or contains(text(),'沒有文章')] | //div[contains(@class,'empty-result')] | //div[contains(@class,'no-results')]"
        )
        if els:
            print("Detected empty result banner:", [el.text for el in els])
        return len(els) > 0
    except Exception as e:
        print("Exception while detecting no-article banner:", e)
        return False


def _results_are_empty(driver, verbose=True) -> bool:
    """
    Return True if all main (top-level, non-dropdown) tab counters show "(0)".
    Returns False if the results-page tab bar is not found (e.g. wrong page or not loaded).
    When verbose, logs each tab/counter for diagnostics.
    """
    try:
        bar = driver.find_element(
            By.XPATH,
            "//ul[contains(@class,'nav-tabs') and contains(@class,'navbar-nav-pub')]"
        )
        # Only non-dropdown tabs
        items = bar.find_elements(By.XPATH, "./li[not(contains(@class,'dropdown'))]")

        total = 0
        zeros = 0
        debug_lines = [] if verbose else None
        for idx, li in enumerate(items):
            s_list = li.find_elements(By.XPATH, "./a/span")
            found = False
            for s in s_list:
                txt = s.text.strip()
                if debug_lines is not None:
                    debug_lines.append(f"Tab idx={idx}, label={li.text!r}, span_text={txt!r}")
                if txt.startswith("(") and txt.endswith(")"):
                    total += 1
                    if txt == "(0)":
                        zeros += 1
                    found = True
                    break
            if not found and debug_lines is not None:
                debug_lines.append(f"Tab idx={idx} had no <a><span> matching (n)!")
        if debug_lines:
            print("\n".join(debug_lines))
            print(f"Tab counter summary: {zeros} of {total} tabs are (0)")
            print(f"Returning from _results_are_empty: {total > 0 and total == zeros}")
        return total > 0 and total == zeros

    except Exception as e:
        if verbose:
            print("Error in _results_are_empty:", e)
        return False



    
def _dump_tab_counters(driver, st):
    try:
        bar = driver.find_element(
            By.XPATH,
            "//ul[contains(@class,'nav-tabs') and contains(@class,'navbar-nav-pub')]"
        )
        items = bar.find_elements(By.XPATH, "./li[not(contains(@class,'dropdown'))]")
        counts = []
        for li in items:
            spans = li.find_elements(By.TAG_NAME, "span")
            for s in spans:
                txt = s.text.strip()
                if txt.startswith("(") and txt.endswith(")"):
                    label = li.text.split("\n")[0].strip()
                    counts.append(f"{label} {txt}")
        st.write("▶️ Top tab counters: " + " | ".join(counts))
    except Exception as e:
        st.warning(f"Could not read tab counters: {e}")


def _debug_tab_bar(driver, st):
    """
    Write the raw outerHTML of the results tab-bar and each child count.
    """
    try:
        bar = driver.find_element(
            By.XPATH,
            "//ul[contains(@class,'nav-tabs') and contains(@class,'navbar-nav-pub')]"
        )
        st.write("🔍 Raw tab-bar HTML:")
        st.code(bar.get_attribute("outerHTML"))
        items = bar.find_elements(By.TAG_NAME, "li")
        counts = []
        for li in items:
            # get the visible number inside parentheses
            txt = li.text.replace("\n", " ").strip()
            counts.append(txt)
        st.write(f"🔢 Parsed tabs: {counts}")
    except Exception as e:
        st.warning(f"Could not debug tab-bar: {e}")


# =============================================================================
# MEDIA/AUTHOR PANEL HELPERS
# =============================================================================

def _capture_results_screenshot(driver, st, reason: str):
    screenshot_dir = (
        os.getenv("WISERS_SCREENSHOT_DIR")
        or os.path.join(".", "artifacts", "screenshots")
    )
    if not driver:
        return
    try:
        inject_cjk_font_css(driver, st_module=st)
        img_bytes = driver.get_screenshot_as_png()
        if st:
            st.image(img_bytes, caption=f"{reason} screenshot")
        try:
            os.makedirs(screenshot_dir, exist_ok=True)
            ts = time.strftime("%Y%m%d_%H%M%S")
            fname = f"{ts}_{reason}.png"
            local_fp = os.path.join(screenshot_dir, fname)
            with open(local_fp, "wb") as f:
                f.write(img_bytes)
        except Exception:
            pass
    except Exception:
        pass


def _results_are_empty_with_banner(driver) -> bool:
    try:
        return _results_are_empty(driver) or _detect_no_article_banner(driver)
    except Exception:
        return False


def _get_home_inputs(key):
    return (HTML_STRUCTURE.get("home", {}).get("inputs", {}) or {}).get(key, [])


def _get_edit_search_selectors(key):
    return (HTML_STRUCTURE.get("edit_search", {}) or {}).get(key, [])


def _get_edit_search_inputs(key):
    return (HTML_STRUCTURE.get("edit_search", {}).get("inputs", {}) or {}).get(key, [])


def _selector_to_by(selector_def):
    by = (selector_def or {}).get("by")
    value = (selector_def or {}).get("value")
    if not by or not value:
        return None, None
    by_map = {
        "css": By.CSS_SELECTOR,
        "xpath": By.XPATH,
        "id": By.ID,
        "name": By.NAME,
    }
    return by_map.get(by), value


def _find_first_visible_input(driver, wait, selector_defs, timeout=6):
    for sel in selector_defs or []:
        by, value = _selector_to_by(sel)
        if not by or not value:
            continue
        try:
            el = WebDriverWait(driver, timeout).until(
                EC.visibility_of_element_located((by, value))
            )
            if el:
                return el
        except Exception:
            continue
    return None


def _is_edit_search_modal_open(driver):
    title_selectors = _get_edit_search_selectors("modal_title")
    for sel in title_selectors:
        by, selector = _selector_to_by(sel)
        if not by or not selector:
            continue
        try:
            els = driver.find_elements(by, selector)
        except Exception:
            els = []
        for el in els:
            try:
                if not el.is_displayed():
                    continue
                txt = (el.text or "").strip()
                if "编辑搜索" in txt or "編輯搜索" in txt:
                    return True
            except Exception:
                continue
    return False


def _close_edit_search_modal(driver, st=None):
    if not _is_edit_search_modal_open(driver):
        return False
    close_selectors = _get_edit_search_selectors("close_button")
    for sel in close_selectors:
        by, selector = _selector_to_by(sel)
        if not by or not selector:
            continue
        try:
            btns = driver.find_elements(by, selector)
        except Exception:
            btns = []
        for btn in btns:
            try:
                if not btn.is_displayed():
                    continue
                driver.execute_script("arguments[0].click();", btn)
                time.sleep(0.6)
                return True
            except Exception:
                continue
    if st:
        st.warning("未能關閉編輯搜索彈窗，將嘗試繼續流程。")
    return False


def _expand_media_author_panel(driver, wait, st=None):
    """
    Best-effort expand the 媒體/作者 panel for visibility.
    """
    state_cfg = HTML_STRUCTURE.get("home", {}).get("media_author_panel_states") or {}
    expanded_selectors = state_cfg.get("expanded") or []
    collapsed_selectors = state_cfg.get("collapsed") or []
    toggle_selectors = HTML_STRUCTURE.get("home", {}).get("media_author_panel_toggles") or []

    for sel in expanded_selectors:
        by, selector = _selector_to_by(sel)
        if not by or not selector:
            continue
        try:
            el = driver.find_element(by, selector)
            if el and el.is_displayed():
                return True
        except Exception:
            continue

    for sel in collapsed_selectors:
        by, selector = _selector_to_by(sel)
        if not by or not selector:
            continue
        try:
            el = driver.find_element(by, selector)
            if el and el.is_displayed():
                driver.execute_script("arguments[0].click();", el)
                time.sleep(0.6)
                break
        except Exception:
            continue

    for sel in toggle_selectors:
        by, selector = _selector_to_by(sel)
        if not by or not selector:
            continue
        try:
            toggle = wait.until(EC.element_to_be_clickable((by, selector)))
            driver.execute_script("arguments[0].click();", toggle)
            time.sleep(0.6)
            return True
        except Exception:
            continue

    for sel in expanded_selectors:
        by, selector = _selector_to_by(sel)
        if not by or not selector:
            continue
        try:
            el = driver.find_element(by, selector)
            if el and el.is_displayed():
                return True
        except Exception:
            continue
    if st:
        st.warning("未能展开『媒體/作者』面板，继续尝试其他方式。")
    return False


def _apply_media_presets(driver, wait, st=None):
    """
    Apply keyword-search media presets: 報刊 / 綜合新聞 / 香港.
    """
    try:
        keep_labels = HTML_STRUCTURE.get("home", {}).get("media_filter_keep_labels") or []
        container_selector = HTML_STRUCTURE.get("home", {}).get("media_filter_container_selector")
        result = set_media_filters_in_panel(
            driver=driver,
            wait=wait,
            st_module=st,
            keep_labels=keep_labels,
            container_selector=container_selector,
        )
        if result:
            return True
    except Exception:
        pass

    if st:
        st.warning("未能套用媒體預設（報刊/綜合新聞/香港），将继续搜索。")
    return False


# =============================================================================
# AUTHOR SEARCH SPECIFIC FUNCTIONS
# =============================================================================

@retry_step
def perform_author_search(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    author_name = kwargs.get('author')
    st = kwargs.get('st_module')
    watchdog = kwargs.get('watchdog')

    if watchdog:
        watchdog.beat()
    _expand_media_author_panel(driver, wait, st)
    _apply_media_presets(driver, wait, st)

    author_input = _find_first_visible_input(
        driver, wait, _get_home_inputs("author"), timeout=8
    )
    if not author_input:
        raise NoSuchElementException("Could not locate author input field.")

    author_input.clear()
    author_input.send_keys(author_name)
    search_button = wait_for_enabled_search_button(driver, timeout=10, st_module=st)
    search_button.click()
    wait_for_results_panel_ready(driver=driver, wait=wait, st_module=st)
    if watchdog:
        watchdog.beat()

@retry_step
def ensure_search_results_ready(**kwargs):
    """
    Wait up to 12 s for either:
      • a clickable headline, **or**
      • the tab-bar to render with all-zero counters.
    Returns True if headlines exist, False if counters are all zero.
    """
    try:
        driver = kwargs["driver"]
        st     = kwargs.get("st_module")
        watchdog = kwargs.get("watchdog")
        if watchdog:
            watchdog.beat()

        # Ensure results panel progress bar/preloader is done before checks
        try:
            wait_for_results_panel_ready(
                driver=driver,
                wait=kwargs.get("wait"),
                st_module=st,
                timeout=20,
            )
            wait_for_ajax_complete(driver, timeout=10)
        except Exception:
            pass

        headline_cond = EC.element_to_be_clickable(
            (By.CSS_SELECTOR, "div.list-group .list-group-item h4 a")
        )

        def _ready(d):
            # Only run Wisers results-page checks when we're on a Wisers page.
            # Avoids NoSuchElementException spam and 503 when driver is on login/other page.
            try:
                url = (d.current_url or "")
                if "wisers" not in url.lower():
                    return False
            except Exception:
                return False
            ready_headline = False
            try:
                ready_headline = headline_cond(d)
            except Exception:
                pass
            ready_empty = False
            try:
                ready_empty = _results_are_empty(d, verbose=False)
            except Exception:
                pass
            ready_no_article = False
            try:
                ready_no_article = _detect_no_article_banner(d)
            except Exception:
                pass
            return ready_headline or ready_empty or ready_no_article

        WebDriverWait(driver, 12).until(_ready)

        empty = _results_are_empty(driver, verbose=True)
        noarticle = _detect_no_article_banner(driver)

        if st:
            if empty:
                st.info("ℹ️ All tab counters are 0 – no articles.")
            elif noarticle:
                st.info("ℹ️ No-article banner detected – no articles.")
            else:
                st.write("✅ Headlines present – results found.")
        if empty or noarticle:
            _capture_results_screenshot(driver, st, "no_results")
        if watchdog:
            watchdog.beat()
        return not (empty or noarticle)

    except TimeoutException:
        print("TimeoutException! Dumping tab bar state for diagnostics:")
        if st:
            st.warning("TimeoutException! Dumping tab bar state for diagnostics:")
        try:
            bar = driver.find_element(
                By.XPATH,
                "//ul[contains(@class,'nav-tabs') and contains(@class,'navbar-nav-pub')]"
            )
            tab_html = bar.get_attribute("outerHTML")
            print("Tab bar HTML:\n", tab_html)
            if st:
                st.code(tab_html, language='html')
        except Exception as ex:
            print("Could not locate tab bar for dumping HTML:", ex)
            if st:
                st.error(f"Could not locate tab bar for dumping HTML: {ex}")
        try:
            main_panel = driver.find_element(By.CSS_SELECTOR, "body")
            html = main_panel.get_attribute("outerHTML")
            print("BODY outerHTML (truncated):\n", html[:4000])
            if st:
                st.code("BODY\n" + html[:4000], language='html')
        except Exception as ex:
            print("Could not get body outerHTML:", ex)
            if st:
                st.error(f"Could not get body outerHTML: {ex}")
        raise



@retry_step
def scrape_hover_popovers(**kwargs):
    """
    Hover over each search result item and collect the popover (hoverbox) content.
    Includes a specific wait for the content preloader to disappear, ensuring
    the full summary is loaded before scraping.

    Expects:
      - driver: Selenium WebDriver
      - wait: WebDriverWait
      - st_module: optional Streamlit module for logging
    """
    driver = kwargs.get("driver")
    wait = kwargs.get("wait")
    st = kwargs.get("st_module")
    logger = kwargs.get("logger")
    watchdog = kwargs.get("watchdog")
    screenshot_dir = (
        kwargs.get("screenshot_dir")
        or os.getenv("WISERS_SCREENSHOT_DIR")
        or os.path.join(".", "artifacts", "screenshots")
    )
    # 兼容不同调用：max_articles / max_items
    max_items = kwargs.get("max_articles", kwargs.get("max_items"))

    previews = []

    try:
        wait_for_results_panel_ready(driver=driver, wait=wait, st_module=st)
        ensure_results_list_visible(driver=driver, wait=wait, st_module=st)
        elements = driver.find_elements(By.CSS_SELECTOR, "span[rel='popover-article']")
        if isinstance(max_items, int) and max_items > 0:
            elements = elements[:max_items]
        if st:
            st.write(f"Found {len(elements)} hoverable items on the page.")
        if st and len(elements) == 0:
            try:
                inject_cjk_font_css(driver, st_module=st)
                img_bytes = driver.get_screenshot_as_png()
                st.image(img_bytes, caption="🔎 搜索结果页（未发现可悬浮条目）")
                try:
                    os.makedirs(screenshot_dir, exist_ok=True)
                    ts = time.strftime("%Y%m%d_%H%M%S")
                    fname = f"{ts}_no_hover_items.png"
                    local_fp = os.path.join(screenshot_dir, fname)
                    with open(local_fp, "wb") as f:
                        f.write(img_bytes)
                    up_logger = logger or (get_logger(st) if st else None)
                    if up_logger and hasattr(up_logger, "upload_screenshot_bytes"):
                        up_logger.upload_screenshot_bytes(img_bytes, filename=fname)
                except Exception:
                    pass
            except Exception as e:
                st.warning(f"搜索结果页截图失败: {e}")

        actions = ActionChains(driver)

        for idx, el in enumerate(elements):
            if watchdog and idx % 8 == 0:
                watchdog.beat()
            title = el.text.strip()

            try:
                # Move mouse over the element to trigger the popover
                actions.move_to_element(el).perform()

                # 1. Wait for the main popover container to appear
                # We need a short wait here just for the container.
                popover_wait = WebDriverWait(driver, 5)
                popover = popover_wait.until(
                    EC.visibility_of_element_located(
                        (By.CSS_SELECTOR, "div.popover.popover-article")
                    )
                )

                # 2. NEW: Wait for the preloader *inside* the popover to disappear.
                #    This is the key to ensuring the content is fully loaded.
                popover_wait.until(
                    EC.invisibility_of_element_located(
                        (By.CSS_SELECTOR, "div.popover.popover-article div.preloader")
                    )
                )
            

                # 3. Now it's safe to scrape the full content
                hover_html = popover.get_attribute("innerHTML")
                hover_text = popover.text.strip()

                previews.append(
                    {
                        "index": idx,
                        "title": title,
                        "hover_html": hover_html,
                        "hover_text": hover_text,
                    }
                )

            except TimeoutException:
                if st:
                    st.warning(
                        f"Popover content for result {idx+1} ('{title[:30]}...') "
                        "did not fully load in time."
                    )
                previews.append(
                    {
                        "index": idx,
                        "title": title,
                        "hover_html": f"<i>(Content failed to load for: {title})</i>",
                        "hover_text": f"(Content failed to load for: {title})",
                        "error": "popover_content_timeout",
                    }
                )
            except Exception as e:
                if st:
                    st.warning(f"Error capturing hover for item {idx+1}: {e}")
                previews.append(
                    {
                        "index": idx,
                        "title": title,
                        "hover_html": f"<i>(Error: {e})</i>",
                        "hover_text": f"(Error: {e})",
                        "error": str(e),
                    }
                )

        return previews

    except Exception as e:
        if st:
            st.error(f"Critical error in scrape_hover_popovers: {e}")
            st.code(traceback.format_exc())
        # Let retry_step handle retries / final failure
        raise


@retry_step
def click_first_result(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    original_window = kwargs.get('original_window')
    st = kwargs.get('st_module')
    watchdog = kwargs.get('watchdog')
    if watchdog:
        watchdog.beat()

    try:
        wait_for_results_panel_ready(driver=driver, wait=wait, st_module=st, timeout=20)
        wait_for_ajax_complete(driver, timeout=10)
        ensure_results_list_visible(driver=driver, wait=wait, st_module=st)
    except Exception:
        pass
    if _results_are_empty_with_banner(driver):
        _capture_results_screenshot(driver, st, "no_results")
        raise Exception("No results to click.")

    # Make sure the headline is still attached each time we click
    def safe_click():
        headline = wait.until(
            EC.element_to_be_clickable(
                (By.CSS_SELECTOR, "div.list-group .list-group-item h4 a")
            )
        )
        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", headline)
        try:
            headline.click()
        except Exception:
            driver.execute_script("arguments[0].click();", headline)

    safe_click()

    # Wait up to 3 s for a new tab
    try:
        WebDriverWait(driver, 3).until(EC.number_of_windows_to_be(2))
        for handle in driver.window_handles:
            if handle != original_window:
                driver.switch_to.window(handle)
                break
        if st:
            st.write("[click_first_result] Opened in new tab.")
        return
    except TimeoutException:
        if st:
            st.warning("[click_first_result] New tab not detected within 3s, retrying click...")
        # If detail page opened in same tab, accept it
        try:
            if driver.find_elements(By.CSS_SELECTOR, "div.article-detail"):
                if st:
                    st.write("[click_first_result] Opened in same tab.")
                return
        except Exception:
            pass
        # Retry click once
        first_link = wait.until(
            EC.element_to_be_clickable(
                (By.CSS_SELECTOR, "div.list-group .list-group-item h4 a")
            )
        )
        try:
            first_link.click()
        except Exception:
            driver.execute_script("arguments[0].click();", first_link)
        # Then wait again for the new tab
        WebDriverWait(driver, 3).until(EC.number_of_windows_to_be(2))
        for handle in driver.window_handles:
            if handle != original_window:
                driver.switch_to.window(handle)
                break
        if st:
            st.write("[click_first_result] Opened in new tab after retry.")



@retry_step
def parse_media_info_for_author(**kwargs):
    subheading_text = kwargs.get('subheading_text')
    author_name = kwargs.get('author_name')
    st = kwargs.get('st_module')
     

    media_part = subheading_text.split('|')[0].strip()
    page_match = re.search(r'([A-Z]\d{2})', media_part)
    if page_match:
        page_number = page_match.group(1)
        media_name_part = media_part[:page_match.start()].strip()
        mapped_name = next((v for k, v in MEDIA_NAME_MAPPINGS.items() if k in media_name_part), media_name_part)
        return f"{mapped_name} {page_number} {author_name}："

@retry_step
def scrape_author_article_content(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    author_name = kwargs.get('author_name')
    st = kwargs.get('st_module')
    watchdog = kwargs.get('watchdog')
    if watchdog:
        watchdog.beat()
     

    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div.article-detail')))
    time.sleep(3)
    title = driver.find_element(By.CSS_SELECTOR, 'h3').text.strip()
    subheading_text = driver.find_element(By.CSS_SELECTOR, 'div.article-subheading').text.strip()
    media_info = parse_media_info_for_author(subheading_text=subheading_text,author_name=author_name,st_module=st)
    paragraphs = [p.text.strip() for p in driver.find_elements(By.CSS_SELECTOR, 'div.description p') if p.text.strip()]
    if paragraphs:
        formatted_first_paragraph = f"{media_info}{paragraphs[0]}"
        full_content = [formatted_first_paragraph] + paragraphs[1:]
        formatted_content_body = '\n\n'.join(full_content)
        final_output = f"{title}\n\n{formatted_content_body}"
    else:
        final_output = title
    return {'title': title, 'content': final_output}

# =============================================================================
# EDITORIAL SCRAPING FUNCTIONS
# =============================================================================

@retry_step
def run_newspaper_editorial_task(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st = kwargs.get('st_module')

    if not _is_edit_search_modal_open(driver):
        edit_btn = None
        try:
            edit_btn = wait.until(
                EC.element_to_be_clickable(
                    (By.XPATH, "//button[contains(.,'編輯搜索') or contains(.,'编辑搜索')]")
                )
            )
        except Exception:
            edit_btn = None
        if edit_btn:
            try:
                edit_btn.click()
            except Exception:
                driver.execute_script("arguments[0].click();", edit_btn)

        for sel in _get_edit_search_selectors("modal_title"):
            by, selector = _selector_to_by(sel)
            if not by or not selector:
                continue
            try:
                wait.until(EC.visibility_of_element_located((by, selector)))
                break
            except Exception:
                continue

    if not _is_edit_search_modal_open(driver) and st:
        st.warning("未能打开『编辑搜索』弹窗，将继续尝试直接设置搜索条件。")

    _expand_media_author_panel(driver, wait, st)
    _apply_media_presets(driver, wait, st)

    column_input = _find_first_visible_input(
        driver, wait, _get_edit_search_inputs("column"), timeout=8
    )
    if not column_input:
        column_input = _find_first_visible_input(
            driver, wait, _get_home_inputs("column"), timeout=5
        )
    if not column_input:
        raise NoSuchElementException("Could not locate column input field.")
    column_input.clear()
    column_input.send_keys("社評 OR editorial")

    search_btn = None
    selectors = [
        (By.CSS_SELECTOR, "button.edit-search-button-track"),
        (By.XPATH, "//div[contains(@class,'modal-footer')]//button[text()='搜索']"),
        (By.XPATH, "//button[normalize-space(text())='搜索']"),
    ]
    for selector_type, selector in selectors:
        try:
            search_btn = wait.until(EC.element_to_be_clickable((selector_type, selector)))
            break
        except TimeoutException:
            continue
    if search_btn:
        search_btn.click()
    else:
        driver.execute_script("""
            var buttons = document.querySelectorAll('div.modal-footer button');
            for (var i = 0; i < buttons.length; i++) {
                if (buttons[i].textContent.trim() === '搜索') {
                    buttons[i].click(); break;
                }
            }""")
    wait_for_results_panel_ready(driver=driver, wait=wait, st_module=st)
    if _is_edit_search_modal_open(driver):
        try:
            wait.until(lambda d: not _is_edit_search_modal_open(d))
        except Exception:
            _close_edit_search_modal(driver, st)

    if wait_for_search_results(driver=driver, wait=wait, st_module=st):
        # NEW: Scroll to load all content, then also wait for AJAX to finish
        scroll_to_load_all_content(driver=driver, st_module=st)
        wait_for_ajax_complete(driver, timeout=10)

        # Now collect all results, with retries
        articles = []
        for retry in range(3):
            results = driver.find_elements(By.CSS_SELECTOR, 'div.list-group-item.no-excerpt')
            if st:
                st.write(f"[Editorial Scrape] Attempt {retry+1}: {len(results)} items found.")
            for result in results:
                try:
                    title = result.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a').text.strip()
                    media_name_raw = result.find_element(By.CSS_SELECTOR, 'small a').text.strip()
                    mapped_name = next((v for k, v in MEDIA_NAME_MAPPINGS.items() if k in media_name_raw), media_name_raw)
                    article = {'media': mapped_name, 'title': title}
                    if article not in articles:
                        articles.append(article)
                except Exception:
                    continue
            if len(articles) > 0:
                break
            time.sleep(2)
        return articles

    return []


@retry_step
def run_scmp_editorial_task(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st = kwargs.get('st_module')

    _expand_media_author_panel(driver, wait, st)
    _apply_media_presets(driver, wait, st)

    author_input = _find_first_visible_input(
        driver, wait, _get_home_inputs("column"), timeout=8
    )
    if not author_input:
        author_input = _find_first_visible_input(
            driver, wait, _get_home_inputs("author"), timeout=5
        )
    if not author_input:
        raise NoSuchElementException("Could not locate column input field.")

    author_input.clear()
    author_input.send_keys("editorial")
    search_button = wait_for_enabled_search_button(driver, timeout=10, st_module=st)
    search_button.click()
    wait_for_results_panel_ready(driver=driver, wait=wait, st_module=st)

    # Scroll and wait for AJAX after search to maximize completeness
    scroll_to_load_all_content(driver=driver, st_module=st)
    wait_for_ajax_complete(driver, timeout=10)

    if wait_for_search_results(driver=driver, wait=wait, st_module=st):
        articles = []
        for retry in range(3):
            results = driver.find_elements(By.CSS_SELECTOR, 'div.list-group-item.no-excerpt')
            if st:
                st.write(f"[SCMP Editorial Scrape] Attempt {retry+1}: {len(results)} items found.")
            for result in results:
                try:
                    title = result.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a').text.strip()
                    media_name_raw = result.find_element(By.CSS_SELECTOR, 'small a').text.strip()
                    mapped_name = next((v for k, v in MEDIA_NAME_MAPPINGS.items() if k in media_name_raw), None)
                    if mapped_name == 'SCMP':
                        article = {'media': 'SCMP', 'title': title}
                        if article not in articles:
                            articles.append(article)
                except Exception:
                    continue
            if len(articles) > 0:
                break
            time.sleep(2)
        return articles
    return []



@retry_step
def create_docx_report(**kwargs):
    author_articles_data = kwargs.get('author_articles_data')
    editorial_data = kwargs.get('editorial_data')
    author_list = kwargs.get('author_list')
    output_path = kwargs.get('output_path')
    st = kwargs.get('st_module')
     
    from docx import Document

    doc = Document()
    doc.add_heading('指定作者社評', level=1)
    doc.add_paragraph()
    for author in author_list:
        article = author_articles_data.get(author)
        title = article['title'] if article else ""
        doc.add_paragraph(f"{author}：{title}")
    doc.add_paragraph()
    is_first_article = True
    for author in author_list:
        article = author_articles_data.get(author)
        if article and article.get('content'):
            if not is_first_article:
                doc.add_paragraph()
            for paragraph_text in article['content'].split('\n\n'):
                doc.add_paragraph(paragraph_text)
            is_first_article = False

    if editorial_data:
        doc.add_page_break()
        doc.add_heading('報章社評', level=1)
        doc.add_paragraph()
        grouped_editorials = defaultdict(list)
        for article in editorial_data:
            grouped_editorials[article['media']].append(article['title'])

        for media, titles in grouped_editorials.items():
            if len(titles) == 1:
                doc.add_paragraph(f"{media}：{titles[0]}")
            else:
                doc.add_paragraph(f"{media}：1. {titles[0]}")
                for i, title in enumerate(titles[1:], start=2):
                    p = doc.add_paragraph()
                    p.add_run(f"\t{i}. {title}")
    
    doc.save(output_path)
    return output_path
