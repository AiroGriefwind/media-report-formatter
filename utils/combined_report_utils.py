import os
import tempfile
from io import BytesIO
from datetime import datetime, timedelta
import pytz

from docx import Document

from utils.document_utils import (
    extract_document_structure,
    rebuild_document_from_structure,
    add_first_page_header,
    add_first_page_footer,
    add_subsequent_pages_header,
    add_subsequent_pages_footer,
)
from utils.wisers_utils import is_hkt_monday


HKT = pytz.timezone("Asia/Hong_Kong")


def _docx_bytes_to_paragraphs(docx_bytes: bytes) -> list:
    doc = Document(BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


def _trim_blank_lines(lines: list) -> list:
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return lines


def extract_keyword_report_body(docx_bytes: bytes, report_title: str) -> list:
    lines = []
    for text in _docx_bytes_to_paragraphs(docx_bytes):
        stripped = text.strip()
        if not stripped:
            lines.append("")
            continue
        if stripped == report_title:
            continue
        if stripped.startswith("日期：") or stripped.startswith("日期:"):
            continue
        if stripped == "（完）":
            continue
        lines.append(text)
    return _trim_blank_lines(lines)


def extract_web_scraping_sections(docx_bytes: bytes) -> tuple:
    author_lines = []
    editorial_lines = []
    mode = None
    for text in _docx_bytes_to_paragraphs(docx_bytes):
        stripped = text.strip()
        if stripped == "指定作者社評":
            mode = "author"
            continue
        if stripped == "報章社評":
            mode = "editorial"
            continue
        if mode == "author":
            author_lines.append(text)
        elif mode == "editorial":
            editorial_lines.append(text)
    return (_trim_blank_lines(editorial_lines), _trim_blank_lines(author_lines))

def _hkt_date_str(days_delta: int = 0) -> str:
    return (datetime.now(HKT).date() + timedelta(days=days_delta)).strftime("%Y%m%d")


def firebase_docx_exists(fb_logger, filename: str, base_folder: str, date_str: str = None) -> bool:
    if not fb_logger or not getattr(fb_logger, "bucket", None):
        return False
    if not date_str:
        date_str = _hkt_date_str()
    remote_path = f"{base_folder}/{date_str}/{filename}"
    try:
        blob = fb_logger.bucket.blob(remote_path)
        return bool(blob.exists())
    except Exception:
        return False


def load_docx_from_firebase_date(fb_logger, filename: str, base_folder: str, date_str: str = None):
    if not fb_logger or not getattr(fb_logger, "bucket", None):
        return None
    if not date_str:
        date_str = _hkt_date_str()
    remote_path = f"{base_folder}/{date_str}/{filename}"
    try:
        blob = fb_logger.bucket.blob(remote_path)
        if blob.exists():
            return blob.download_as_bytes()
    except Exception:
        return None
    return None


def merge_keyword_report_bodies(docx_bytes_list: list, report_title: str) -> list:
    combined = []
    for docx_bytes in docx_bytes_list:
        if not docx_bytes:
            continue
        combined.extend(extract_keyword_report_body(docx_bytes, report_title))
    return _trim_blank_lines(combined)


def build_combined_report_docx_bytes(
    editorial_lines: list,
    international_lines: list,
    greater_china_lines: list,
    local_lines: list,
    author_lines: list,
) -> bytes:
    doc = Document()

    doc.add_paragraph("報章社評")
    for line in editorial_lines:
        doc.add_paragraph(line)

    doc.add_paragraph("國際新聞")
    for line in international_lines:
        doc.add_paragraph(line)

    doc.add_paragraph("大中華新聞")
    for line in greater_china_lines:
        doc.add_paragraph(line)

    doc.add_paragraph("本地新聞")
    for line in local_lines:
        doc.add_paragraph(line)
    for line in author_lines:
        doc.add_paragraph(line)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    try:
        with open(tmp_path, "rb") as f:
            return f.read()
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass


def format_docx_bytes_with_workflow(docx_bytes: bytes) -> bytes:
    monday_mode = is_hkt_monday()
    sunday_date = None
    if monday_mode:
        sunday_date = (datetime.now(HKT).date() - timedelta(days=1)).strftime("%Y%m%d")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
        tmp_in.write(docx_bytes)
        input_path = tmp_in.name

    output_path = input_path.replace(".docx", "_formatted.docx")

    try:
        extract_document_structure(input_path, monday_mode=monday_mode, sunday_date=sunday_date)
        rebuilt_path = rebuild_document_from_structure(
            input_path, monday_mode=monday_mode, sunday_date=sunday_date
        )

        logo_path = os.path.join("assets", "AsiaNet_logo.png")
        doc = Document(rebuilt_path)
        if os.path.exists(logo_path):
            add_first_page_header(doc, logo_path)
        else:
            add_first_page_header(doc, None)
        add_first_page_footer(doc)
        add_subsequent_pages_header(doc)
        add_subsequent_pages_footer(doc)
        doc.save(output_path)

        with open(output_path, "rb") as f:
            return f.read()
    finally:
        for path in [input_path, output_path]:
            try:
                if path and os.path.exists(path):
                    os.remove(path)
            except Exception:
                pass
