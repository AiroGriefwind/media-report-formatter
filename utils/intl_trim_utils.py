#!/usr/bin/env python3
"""trim_from_json_any_order.py

根据 user_final_list.json 中的 title（不依赖 JSON 顺序）定位 docx 里的每篇文章标题，
并输出：标题 + metadata(标题后第一条非空段落) + 正文前三段（非空段落）。

用法：
  python trim_from_json_any_order.py <input.docx> <user_final_list.json> [output.docx]
"""

import os, sys, json, re
from collections import defaultdict
from docx import Document

def is_subtitle_candidate(text, prev_text, next_text, debug=False):
    """ 
    副标题识别：
    - 优先沿用原逻辑：前后为空行 + 短句 + 不以句号结尾。
    - 增强逻辑（适配 docx 中“没有空行，但视觉上是小标题”的情况）：
      若该行很短且前后两段都很长，则视为副标题。

    返回：bool；debug=True 时同时返回 (bool, reason)
    """
    if not text or not text.strip():
        return (False, 'empty') if debug else False

    t = text.strip()

    # Hardcode ignore for section headers
    if t in ["國際新聞", "大中華新聞", "本地新聞"]:
        return (False, 'section_header') if debug else False

    # 常见非正文标记：避免误判
    if t in ["####", "（完）"]:
        return (False, 'marker') if debug else False

    # 必须是短行（副标题通常很短）
    if len(t) > 20:
        return (False, 'too_long') if debug else False

    # 副标题一般不以句号结尾
    if t.endswith('。') or t.endswith('.'):
        return (False, 'ends_with_period') if debug else False

    prev_is_blank = (not prev_text) or (not prev_text.strip())
    next_is_blank = (not next_text) or (not next_text.strip())

    # 原规则：前后都是空行
    if prev_is_blank and next_is_blank:
        return (True, 'blank_surrounded') if debug else True

    # 增强规则：docx 常见情况——没有空行，但“夹在两段长正文之间”的短行
    prev_len = len(prev_text.strip()) if prev_text else 0
    next_len = len(next_text.strip()) if next_text else 0

    if prev_len >= 40 and next_len >= 40:
        return (True, f'between_long_paras(prev={prev_len},next={next_len})') if debug else True

    return (False, f'no_rule_match(prev_blank={prev_is_blank},next_blank={next_is_blank},prev={prev_len},next={next_len})') if debug else False


def copy_paragraph(dst_doc, src_para):
    p = dst_doc.add_paragraph()
    for r in src_para.runs:
        nr = p.add_run(r.text)
        nr.bold = r.bold
        nr.italic = r.italic
        nr.underline = r.underline
        if r.font.size:
            nr.font.size = r.font.size
        if r.font.name:
            nr.font.name = r.font.name
    p.style = src_para.style
    pf_src = src_para.paragraph_format
    pf_dst = p.paragraph_format
    pf_dst.left_indent = pf_src.left_indent
    pf_dst.right_indent = pf_src.right_indent
    pf_dst.first_line_indent = pf_src.first_line_indent
    pf_dst.space_before = pf_src.space_before
    pf_dst.space_after = pf_src.space_after
    pf_dst.line_spacing = pf_src.line_spacing
    pf_dst.keep_with_next = pf_src.keep_with_next
    pf_dst.keep_together = pf_src.keep_together
    pf_dst.alignment = pf_src.alignment
    return p


def normalize_title(text: str) -> str:
    if not text:
        return ""
    t = text.strip()
    # 去掉零宽字符、NBSP
    t = t.replace("\u200b", "").replace("\xa0", " ")
    # 去掉 Markdown **
    t = re.sub(r'^\*{1,2}', '', t)
    t = re.sub(r'\*{1,2}$', '', t)
    # 去掉编号 1. / 1、 / (1)
    t = re.sub(r'^\s*[\(（]?\d+[\)）]?[\.、:]?\s*', '', t)
    # 统一空白
    t = re.sub(r'\s+', ' ', t)
    return t.strip()


def load_titles(json_path: str):
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    titles = []
    for _, items in data.items():
        if not isinstance(items, list):
            continue
        for it in items:
            if isinstance(it, dict) and it.get('title'):
                titles.append(normalize_title(it['title']))
    # 去重但保序
    seen = set()
    uniq = []
    for t in titles:
        if t and t not in seen:
            uniq.append(t)
            seen.add(t)
    return uniq


def build_doc_title_index(doc: Document):
    """扫描整个 doc，把每个段落的 normalized 文本映射到索引列表。"""
    mp = defaultdict(list)
    for i, p in enumerate(doc.paragraphs):
        nt = normalize_title(p.text)
        if nt:
            mp[nt].append(i)
    return mp


def pick_title_indices(doc, titles):
    """为每个 title 从 doc 中挑一个索引（若重复则按出现顺序逐个取）。"""
    mp = build_doc_title_index(doc)
    used = defaultdict(int)
    idxs = []
    missing = []

    for t in titles:
        if t not in mp:
            # 允许“标题只匹配到 doc 段落的子串”的情况：做一次兜底模糊匹配
            candidates = []
            for k, lst in mp.items():
                if (t in k) or (k in t):
                    candidates.append((abs(len(k) - len(t)), k, lst))
            candidates.sort(key=lambda x: x[0])
            if candidates:
                _, k_best, lst_best = candidates[0]
                j = used[k_best]
                if j < len(lst_best):
                    idxs.append(lst_best[j])
                    used[k_best] += 1
                    continue
            missing.append(t)
            continue

        j = used[t]
        if j >= len(mp[t]):
            missing.append(t)
            continue
        idxs.append(mp[t][j])
        used[t] += 1

    if missing:
        # 打印前几个 missing，方便定位
        raise ValueError('以下标题在 docx 找不到（或重复次数不够）：\n- ' + '\n- '.join(missing[:10]))

    # 关键：按 doc 中出现顺序排序，完全不依赖 JSON 顺序
    return sorted(set(idxs))


def trim_docx(input_docx, json_path, output_docx, keep_body_paras=3):
    doc = Document(input_docx)
    new_doc = Document()

    titles = load_titles(json_path)
    title_idxs = pick_title_indices(doc, titles)
    title_set = set(title_idxs)

    i = 0
    while i < len(doc.paragraphs):
        if i in title_set:
            # 1) 标题
            copy_paragraph(new_doc, doc.paragraphs[i])
            i += 1

            # 2) metadata：标题后第一条非空段落（不计入正文段数）
            while i < len(doc.paragraphs) and not doc.paragraphs[i].text.strip():
                i += 1
            if i < len(doc.paragraphs) and i not in title_set:
                copy_paragraph(new_doc, doc.paragraphs[i])
                i += 1

            # 3) 正文前三段（非空），跳过副标题且不计数
            kept = 0
            while i < len(doc.paragraphs) and i not in title_set:
                cur = doc.paragraphs[i].text
                if cur and cur.strip():
                    prev_txt = doc.paragraphs[i-1].text if i > 0 else ""
                    next_txt = doc.paragraphs[i+1].text if i + 1 < len(doc.paragraphs) else ""

                    # 副标题：跳过，不输出，也不计入 kept
                    if is_subtitle_candidate(cur, prev_txt, next_txt):
                        i += 1
                        continue

                    if kept < keep_body_paras:
                        copy_paragraph(new_doc, doc.paragraphs[i])
                        kept += 1
                i += 1

            # 关键：处理完一篇文章后，直接进入下一轮，让下一篇标题被正确识别
            continue

        # 文章之外内容：原样保留（日期、分隔符、（完）等）
        copy_paragraph(new_doc, doc.paragraphs[i])
        i += 1

    new_doc.save(output_docx)
    return output_docx


# if __name__ == '__main__':
#     # 支持：python script.py input.docx user_final_list.json [output.docx] [--debug-subtitle]
#     args = sys.argv[1:]
#     debug_subtitle = False
#     if '--debug-subtitle' in args:
#         debug_subtitle = True
#         args = [a for a in args if a != '--debug-subtitle']

#     if len(args) < 2:
#         print('Usage: python trim_from_json_any_order_fixed_v2.py <input.docx> <user_final_list.json> [output.docx] [--debug-subtitle]')
#         sys.exit(1)

#     input_docx = args[0]
#     json_path = args[1]
#     output_docx = args[2] if len(args) >= 3 else os.path.splitext(input_docx)[0] + '_trimmed.docx'

#     if not os.path.exists(input_docx):
#         raise SystemExit(f'File not found: {input_docx}')
#     if not os.path.exists(json_path):
#         raise SystemExit(f'File not found: {json_path}')

#     # 若开启 debug，会在保留正文时打印每个“短行”的判断原因
#     if not debug_subtitle:
#         out = trim_docx(input_docx, json_path, output_docx)
#         print('OK:', out)
#     else:
#         # debug 模式：复用 trim_docx 逻辑，但在正文扫描时打印潜在副标题判断
#         doc = Document(input_docx)
#         titles = load_titles(json_path)
#         title_idxs = pick_title_indices(doc, titles)
#         title_set = set(title_idxs)

#         i = 0
#         while i < len(doc.paragraphs):
#             if i in title_set:
#                 title_text = doc.paragraphs[i].text.strip()
#                 i += 1
#                 # metadata
#                 while i < len(doc.paragraphs) and not doc.paragraphs[i].text.strip():
#                     i += 1
#                 if i < len(doc.paragraphs) and i not in title_set:
#                     i += 1

#                 # scan body until next title, only debug short lines
#                 while i < len(doc.paragraphs) and i not in title_set:
#                     cur = doc.paragraphs[i].text
#                     if cur and cur.strip() and len(cur.strip()) <= 20:
#                         prev_txt = doc.paragraphs[i-1].text if i > 0 else ""
#                         next_txt = doc.paragraphs[i+1].text if i + 1 < len(doc.paragraphs) else ""
#                         ok, reason = is_subtitle_candidate(cur, prev_txt, next_txt, debug=True)
#                         print(f"[subtitle_debug] idx={i} ok={ok} reason={reason} | article={title_text} | text={cur.strip()}")
#                     i += 1
#                 continue
#             i += 1

#         print('Debug done. Now run without --debug-subtitle to generate output.')
