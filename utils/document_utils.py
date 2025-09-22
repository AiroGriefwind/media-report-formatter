# =============================================================================
# DOCUMENT PROCESSING FUNCTIONS
# =============================================================================
import streamlit as st
import os
import re
import json
from datetime import datetime
from collections import defaultdict
import time

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from opencc import OpenCC

from .config import CORRECTION_MAP, MEDIA_NAME_MAPPINGS, EDITORIAL_MEDIA_ORDER, EDITORIAL_MEDIA_NAMES, TITLE_MODIFICATIONS


# =============================================================================
# DOCUMENT FORMATTING FUNCTIONS
# =============================================================================

def sanitize_doc_paragraphs(doc):
    """
    Clean all paragraphs to plain text (like ctrl+shift+v pasted).
    Removes all formatting, replaces all whitespace variants with normal spaces,
    and collapses multiple linebreaks or non-printing chars.
    """
    sanitized = []
    for para in doc.paragraphs:
        # Unify all whitespace chars to a space/newline
        text = para.text
        text = text.replace('\u200b', '').replace('\xa0', ' ')  # Remove invisible non-breaking spaces etc.
        # Replace any other suspicious chars as you encounter
        # Replace all \r and \n with single newlines
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Collapse multiple spaces
        text = re.sub(r'[ \t]+', ' ', text)
        # Collapse multiple newlines
        text = re.sub(r'\n+', '\n', text)
        # Strip leading/trailing whitespace
        text = text.strip()
        sanitized.append(text)
    return sanitized


def is_source_citation(text):
    """Check if text is a source citation"""
    if not text: 
        return False
    if ']' in text and text.index(']') < 30: 
        return False
    if re.match(r'^.{1,20}[:：]', text): 
        return True
    common_media_prefixes = "|".join(re.escape(k) for k in MEDIA_NAME_MAPPINGS.keys())
    if re.match(rf'^({common_media_prefixes})\s*[:：]', text): 
        return True
    return False

def is_valid_headline(text):
    """Validates if a line of text could be a headline"""
    if not text or len(text.strip()) < 5:
        return False
    if re.search(r'[，,。]', text):
        return False
    if re.search(r'[.?!]$', text.strip()):
        return False
    if ']' in text:
        return False
    return True

def is_new_metadata_format(text):
    """
    Detects if a line is in the new metadata format.
    Format: "media name + page number + section name | word count | date"
    """
    if not text:
        return False
    
    # Check for exactly two '|' characters
    if text.count('|') != 2:
        return False
    
    # Check that it doesn't end with punctuation
    if re.search(r'[。，.，]$', text.strip()):
        return False
    
    # Additional validation - should have format like "媒体名 页码 栏目名 |字数 |日期"
    parts = text.split('|')
    if len(parts) != 3:
        return False
    
    # The first part should contain media name and page info
    first_part = parts[0].strip()
    if not first_part:
        return False
    
    # The second part should contain word count (should have '字')
    second_part = parts[1].strip()
    if '字' not in second_part:
        return False
    
    # The third part should look like a date
    third_part = parts[2].strip()
    if not re.match(r'^\d{4}-\d{2}-\d{2}$', third_part):
        return False
    
    return True

def transform_metadata_line(metadata_text, next_paragraph_text):
    if not metadata_text or not next_paragraph_text:
        return metadata_text

    # 檢查多份報章專用標記
    has_placeholder = "==" in metadata_text

    # Split by |
    parts = metadata_text.split('|')
    if len(parts) < 1:
        return metadata_text

    main_part = parts[0].replace('==', '').strip()  # 徹底移除 '=='
    tokens = main_part.split()
    # 安全檢查
    if len(tokens) >= 2:
        media_name = tokens[0]
        page_number = tokens[1]
    else:
        media_name = tokens[0]
        page_number = ''

    body = remove_reporter_phrases(next_paragraph_text.strip())
    short_media_name = get_short_media_name(media_name)

    # 組裝：及多份報章加在page_number後
    page_label = f"{page_number}{'及多份報章' if has_placeholder and page_number else ''}"
    # 如沒頁碼直接加在媒體名後（極少見，保底）
    if has_placeholder and not page_number:
        short_media_name += '及多份報章'

    transformed = f"{short_media_name} {page_label}：{body}".replace('  ', ' ').strip()
    return transformed



def get_short_media_name(full_media_name):
    """
    Get short media name from full media name using MEDIA_NAME_MAPPINGS.
    """
    # Import here to avoid circular dependency
    from .config import MEDIA_NAME_MAPPINGS

    # Direct match
    if full_media_name in MEDIA_NAME_MAPPINGS:
        return MEDIA_NAME_MAPPINGS[full_media_name]
    # Try first word match
    first = full_media_name.split()[0]
    if first in MEDIA_NAME_MAPPINGS:
        return MEDIA_NAME_MAPPINGS[first]
    # Fallback: try longest matching prefix
    return first



def is_subtitle_candidate(text, prev_text, next_text):
    """
    Detect if a line is a subtitle that should be removed.
    Criteria:
    1. Has blank line before and after
    2. Never ends with period (。 or .)
    3. Never exceeds 20 Chinese characters
    4. Not empty
    5. Not a section header
    """
    if not text or not text.strip():
        return False
    
    text = text.strip()
    
    # Hardcode ignore for section headers
    if text in ["國際新聞", "大中華新聞", "本地新聞"]:
        return False
    
    # Check if previous and next texts are blank/empty
    prev_is_blank = not prev_text or not prev_text.strip()
    next_is_blank = not next_text or not next_text.strip()
    
    if not (prev_is_blank and next_is_blank):
        return False
    
    # Check if it ends with a period
    if text.endswith('。') or text.endswith('.'):
        return False
    
    # Check character count (20 or less Chinese characters)
    if len(text) > 20:
        return False
    
    return True

def remove_inline_figure_table_markers(text: str) -> str:
    """
    Remove inline references to figures/tables that appear inside parentheses/brackets.
    Handles cases like:
      - （見圖）, （见表）, （详见表）, （圖一）, （图1）, 【圖】, 【见表】, [G1]
      - Mixed content where a trailing marker follows valid text:
          e.g. "（Terence Tao，小圖[G1]）" -> "（Terence Tao）"
    Works on full-width and ASCII parentheses/brackets, in Simplified/Traditional.
    """
    if not text:
        return text

    import re
    num_cn = '一二三四五六七八九十百千零〇两兩'
    markers_prefix = r'(?:見|见|详见|詳見|小|附)?'
    marker_core = r'(?:图|圖|表)'
    marker_suffix = rf'(?:\s*(?:[0-9]+|[{num_cn}]{{1,3}}))?'
    full_marker = markers_prefix + r'\s*' + marker_core + marker_suffix

    # Precompile regexes
    re_sq_bracket_token = re.compile(r'\[\s*[GPgp]\d+\s*\]')
    re_trailing_marker = re.compile(rf'[，,、]\s*{full_marker}\s*$')
    re_full_marker_only = re.compile(rf'^\s*{full_marker}\s*$')
    re_sq_or_fw_marker_only = re.compile(rf'^\s*(?:{full_marker}|[GPgp]\d+)\s*$')

    s = text

    # 1) Remove simple [G1]/[P2] tokens anywhere
    s = re_sq_bracket_token.sub('', s)

    # 2) Clean parentheses content while preserving earlier meaningful parts
    for name, rx in [('fullwidth', re.compile(r'（([^（）]*)）')),
                     ('ascii', re.compile(r'\(([^()]*)\)'))]:
        def paren_repl(m):
            inner = m.group(1) or ''
            # Remove any residual [G1]-like tokens inside
            inner_clean = re_sq_bracket_token.sub('', inner)
            # If the parentheses contain only a marker like (见图1)/(表三)/(小图) → drop entire ()
            if re_full_marker_only.match(inner_clean.strip()):
                return ''
            # If marker trails after a comma, drop that trailing segment
            kept = re_trailing_marker.sub('', inner_clean).strip()
            if not kept:
                return ''
            return ('（' if name == 'fullwidth' else '(') + kept + ('）' if name == 'fullwidth' else ')')
        s = rx.sub(paren_repl, s)

    # 3) Remove bracketed segments that are purely markers (【见表】, [见图2], 【图】)
    for name, rx in [('square', re.compile(r'\[([^\[\]]*)\]')),
                     ('fullwidth', re.compile(r'【([^【】]*)】'))]:
        def bracket_repl(m):
            inner = (m.group(1) or '').strip()
            if re_sq_or_fw_marker_only.match(inner):
                return ''
            return m.group(0)
        s = rx.sub(bracket_repl, s)

    # 4) Cleanup empty pairs and spacing
    s = re.sub(r'\(\s*\)', '', s)
    s = re.sub(r'（\s*）', '', s)
    s = re.sub(r'\[\s*\]', '', s)
    s = re.sub(r'【\s*】', '', s)
    s = re.sub(r'\s{2,}', ' ', s)
    s = re.sub(r'\s+([，。；：、？！)])', r'\1', s)
    s = re.sub(r'（\s+', '（', s)

    return s.strip()


def remove_reporter_phrases(text):
    if not text:
        return ""
    #Remove `●香港文匯報記者 or 香港文汇报记者` and anything after it
    text = re.sub(r'(●香港文匯報記者|●香港文汇报记者|大公文匯全媒體記者|大公文汇全媒体记者).*$', '', text, flags=re.MULTILINE)

    #Remove reporting agency content between first colon and '报道：' or '報道：'
    match = re.search(r'(^.+?：)', text)
    if match:
        prefix = match.group(1)  # everything up to and including the first colon
        rest = text[match.end():]  # everything after the first colon
        
        # Find 报道： or 報道： in the rest
        rep_match = re.search(r'(报道：|報道：)', rest)
        if rep_match:
            # Remove text from start of rest up to and including 报道：/報道：
            rest = rest[rep_match.end():].lstrip()
            text = prefix + rest

    # Remove 【...】 containing keywords
    pattern_brackets = r'【[^】]*?(记者|記者|报道|報道|报讯|報訊|專訊|专讯)[^】]*?】'
    text = re.sub(pattern_brackets, '', text)
    # Remove （...） containing keywords, using a function for precision
    def paren_replacer(match):
        if re.search(r'(记者|記者|报道|報道|报讯|報訊|專訊|专讯)', match.group(1)):
            return ''
        return match.group(0)
    text = re.sub(r'（([^）]*)）', paren_replacer, text)
    # Remove the fixed phrase
    text = text.replace('香港文汇报訊', '').replace('香港文匯報訊', '')
    return text.strip()

def convert_to_traditional_chinese(text):
    """Convert simplified Chinese to traditional Chinese"""
    if not text or not text.strip():
        return text
    try:
        cc = OpenCC('s2hk')
        converted_text = cc.convert(text)
        return converted_text
    except Exception as e:
        print(f"Warning: Chinese conversion failed for text: {text[:50]}... Error: {str(e)}")
        return text
    
def apply_gatekeeper_corrections(text):
    """
    Applies a second round of specific corrections based on the CORRECTION_MAP.
    """
    if not text:
        return text
    for error, correction in CORRECTION_MAP.items():
        if error in text:
            text = text.replace(error, correction)
    return text

def add_end_marker(doc):
    """Add end marker to document"""
    blank_para = doc.add_paragraph("")
    blank_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    end_para = doc.add_paragraph("（完）")
    end_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    end_para.style = doc.styles['Normal']

def setup_document_fonts(doc):
    """Setup document fonts"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def add_first_page_header(doc, logo_path):
    """Add header only on the first page"""
    try:
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        header = section.first_page_header
        
        # Always ensure there's at least one paragraph
        if header.paragraphs:
            header_para = header.paragraphs[0]
            header_para.clear()
        else:
            header_para = header.add_paragraph()

        # Add the text run (always)
        left_run = header_para.add_run("亞聯每日報章摘要")
        left_run.font.name = 'Calibri'
        left_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        left_run.font.size = Pt(18)

        # Add tab stops and logo (only if logo exists)
        if logo_path and os.path.exists(logo_path):
            tab_stops = header_para.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            header_para.add_run("\t\t")
            logo_run = header_para.add_run()
            logo_run.add_picture(logo_path, width=Cm(5.95), height=Cm(2.04))

        header_para.style = doc.styles['Header']
        
    except Exception as e:
        print(f"Warning: Could not add first page header: {e}")



def add_first_page_footer(doc):
    """Add footer only on the first page"""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    footer = section.first_page_footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    
    footer_lines = [
        "香港金鐘夏愨道18號海富中心24樓  電話: 2114 4960  傳真: 3544 2933",
        "電郵: info@asianet-sprg.com.hk", 
        "網頁: http://www.asianet-sprg.com.hk"
    ]
    
    for i, line in enumerate(footer_lines):
        run = footer_para.add_run(line)
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(12)
        if i < len(footer_lines) - 1:
            footer_para.add_run('\n')
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_subsequent_pages_header(doc):
    """Add header to pages 2 onwards"""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.clear()
    
    header_run = header_para.add_run("AsiaNet亞聯政經顧問")
    header_run.font.name = 'Calibri'
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    header_run.font.size = Pt(12)
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_subsequent_pages_footer(doc):
    """Add footer to pages 2 onwards with page numbers"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    
    footer_run = footer_para.add_run()
    footer_run.font.name = 'Calibri'
    footer_run.font.size = Pt(12)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    footer_run._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    footer_run._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    footer_run._element.append(fldChar2)
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def detect_section_type(text):
    """Detect the type of section from text"""
    if not text: 
        return None
    sections = {
        'editorial': r'^報章社評\s*$', 
        'international': r'^國際新聞[:：]?\s*$', 
        'china': r'^大中華新聞\s*$', 
        'local': r'^本地新聞\s*$', 
        'financial': r'^財經新聞\s*$', 
        'Hong Kong': r'^香港本地新聞\s*$', 
        'entertainment': r'^娛樂新聞\s*$', 
        'sports': r'^體育新聞\s*$', 
        'property': r'^地產新聞\s*$'
    }
    for name, pattern in sections.items():
        if re.match(pattern, text): 
            return name
    return None

def detect_editorial_media_line(text):
    """Detects editorial media lines"""
    if not text: 
        return None
    
    match = re.match(r'^([^：]+)：(.*)$', text)
    if match:
        potential_name, content = match.group(1).strip(), match.group(2).strip()
        
        if potential_name in EDITORIAL_MEDIA_NAMES:
            return {'full_name': potential_name, 'clean_name': potential_name, 'content': content}
        
        if potential_name in MEDIA_NAME_MAPPINGS:
            clean_name = MEDIA_NAME_MAPPINGS[potential_name]
            return {'full_name': potential_name, 'clean_name': clean_name, 'content': content}
    
    return None

def is_editorial_continuation(text):
    """Detects if a line is a continuation of editorial content"""
    if not text: 
        return False
    if re.match(r'^\s*\d+\.\s+', text): 
        return True
    if re.match(r'^[\t\s]{2,}', text): 
        return True
    if len(text.strip()) > 15: 
        return True
    return False

def format_content_paragraph(paragraph):
    """Format content paragraph"""
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

def format_media_first_line_hanging(paragraph, label_length):
    """Format media first line with hanging indent"""
    pf = paragraph.paragraph_format
    indent_amount = Pt(54)
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.keep_with_next = True
    pf.left_indent = indent_amount
    pf.first_line_indent = -indent_amount
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(indent_amount)

def format_section_header(paragraph):
    """Format section header"""
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.left_indent = Pt(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True

def format_article_title(paragraph, needs_spacing):
    """Format article title"""
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(12) if needs_spacing else Pt(0)
    pf.space_after = Pt(0)
    pf.keep_with_next = True

def add_section_header_to_doc(doc, text):
    """Add section header to document"""
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    format_section_header(p)

def add_article_to_document(doc, article_data, needs_spacing):
    """Add article to document"""
    p = doc.add_paragraph(article_data['text'], style='List Number')
    p.runs[0].bold = True
    #p.style = doc.styles['Normal']
    format_article_title(p, needs_spacing)

def add_media_group_to_document(new_doc, media_group):
    """Add media group to document"""
    media_label = f"{media_group['clean_name']}："
    label_length = len(media_label)
    full_width_space = '\u3000'

    para = new_doc.add_paragraph()
    para.add_run(f"{media_label}{media_group['first_item']}")
    format_media_first_line_hanging(para, label_length)

    for item in media_group['additional_items']:
        item_para = new_doc.add_paragraph()
        item_para.add_run(full_width_space * label_length + item['text'])
        format_media_first_line_hanging(item_para, label_length)

def add_monday_notice(new_doc, sunday_date):
    """Add Monday notice to document"""
    notice_line = f"是日新聞摘要包括週日重點新聞，除註明{sunday_date}外，其他均是今天新聞"
    new_doc.add_paragraph(notice_line)

def add_date_line_if_needed(doc, date_str):
    """
    Add date line if needed, safely handling empty documents.
    This corrected version avoids the index error.
    """
    # First, check if the document already has a date line.
    if doc.paragraphs:
        first_text_paragraph = next((p for p in doc.paragraphs if p.text.strip()), None)
        if first_text_paragraph and re.match(r'^\d{8}$', first_text_paragraph.text.strip()):
            return  # Date already exists, so we do nothing.

    # If the document is empty or the date is missing, we add it.
    # This method safely adds the paragraph and moves it to the beginning.
    p = doc.add_paragraph(date_str)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_element = p._p
    doc._body._element.remove(p_element)
    doc._body._element.insert(0, p_element)


def add_end_marker(doc):
    """Add end marker to document"""
    blank_para = doc.add_paragraph("")
    blank_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    end_para = doc.add_paragraph("（完）")
    end_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    end_para.style = doc.styles['Normal']

# =============================================================================
# MAIN DOCUMENT PROCESSING FUNCTIONS
# =============================================================================

def extract_document_structure(doc_path, json_output_path=None, is_monday_mode=False, sunday_date=None):
    """
    Extracts structure using state-based logic with Chinese conversion.
    """
    global TITLE_MODIFICATIONS
    TITLE_MODIFICATIONS = []

    if json_output_path is None:
        json_output_path = doc_path.replace('.docx', '_structure.json')
    
    doc = Document(doc_path)
    sanitized_paragraphs = sanitize_doc_paragraphs(doc)
    structure = {'total_paragraphs': len(sanitized_paragraphs), 'editorial_media_groups': [], 'sections': {}, 'other_content': []}
    
    current_section = None
    in_editorial = False
    section_counters = {}
    
    is_expecting_title = False
    title_cooldown_counter = 0

    num_paragraphs = len(sanitized_paragraphs)
    # add paragraph while skipping first paragraph after metadata
    skip_next = False
    for i, paragraph in enumerate(sanitized_paragraphs):
        if skip_next:
            skip_next = False
            continue  # Skip the paragraph after the metadata line

        
        original_text = paragraph.strip()
        text = convert_to_traditional_chinese(original_text)
        text = apply_gatekeeper_corrections(text)

        text = remove_reporter_phrases(text)
        text = remove_inline_figure_table_markers(text)

         # Check for subtitle removal BEFORE other processing
        if i > 0 and i < num_paragraphs - 1:  # Not first or last paragraph
            prev_text = sanitized_paragraphs[i-1].strip() if i > 0 else ""
            next_text = sanitized_paragraphs[i+1].strip() if i < num_paragraphs - 1 else ""
            
            # Convert previous and next text for proper comparison
            prev_text = convert_to_traditional_chinese(prev_text)
            next_text = convert_to_traditional_chinese(next_text)
            
            import streamlit as st
            if is_subtitle_candidate(text, prev_text, next_text):
                # This is a subtitle, skip it
                st.write("subtitle found and removed:", text)
                structure['other_content'].append({
                    'index': i, 
                    'text': text, 
                    'type': 'subtitle_removed', 
                    'section': current_section
                })
                continue  # Skip this subtitle

        if is_new_metadata_format(original_text):
            next_content = ""
            if i + 1 < num_paragraphs:
                next_paragraph_text = sanitized_paragraphs[i + 1].strip()
                next_content = convert_to_traditional_chinese(next_paragraph_text)
                next_content = apply_gatekeeper_corrections(next_content)
            
            # Extract date for Sunday detection
            date_str = original_text.split('|')[-1].strip().replace("-", "")
            is_sunday_article = is_monday_mode and (date_str == sunday_date)
            
            text = transform_metadata_line(text, next_content)
            skip_next = True

            # Now when creating media info or article, store this flag
            # For example:
            media_info = detect_editorial_media_line(text)
            if media_info:
                # create dict for current_media_group with is_sunday_article stored
                current_media_group = {
                    'clean_name': media_info['clean_name'], 
                    'original_name': media_info['full_name'], 
                    'start_index': i, 
                    'first_item': next_content, 
                    'additional_items': [],
                    'is_sunday_article': is_sunday_article  # store flag here
                }
                structure['editorial_media_groups'].append(current_media_group)
            continue
        
        if not text: continue
        
        if in_editorial:
            media_info = detect_editorial_media_line(text)
            if media_info:
                # Apply conversion to media content as well
                converted_content = convert_to_traditional_chinese(media_info['content'])
                # Apply gatekeeper corrections to media content
                converted_content = apply_gatekeeper_corrections(converted_content)
                # Remove reporter phrases from media content
                converted_content = remove_reporter_phrases(converted_content)
                # Remove inline figure/table markers from media content
                converted_content = remove_inline_figure_table_markers(converted_content)
                

                media_info['content'] = converted_content
                current_media_group = {'clean_name': media_info['clean_name'], 'original_name': media_info['full_name'], 'start_index': i, 'first_item': converted_content, 'additional_items': []}
                structure['editorial_media_groups'].append(current_media_group)
            elif 'current_media_group' in locals() and current_media_group and is_editorial_continuation(text):
                current_media_group['additional_items'].append({'index': i, 'text': text})
        else:
            # Rest of the existing logic remains the same, but with converted text
            if title_cooldown_counter > 0:
                structure['other_content'].append({'index': i, 'text': text, 'type': 'content', 'section': current_section})
                title_cooldown_counter -= 1
                continue

            is_title = False
            prospective_title_text = text

            if is_expecting_title:
                is_title = True
                is_expecting_title = False
            else:
                if i + 1 < num_paragraphs:
                    next_paragraph_text_original = sanitized_paragraphs[i+1].strip()
                    next_paragraph_text = convert_to_traditional_chinese(next_paragraph_text_original)
                    next_paragraph_text = apply_gatekeeper_corrections(next_paragraph_text)
                    next_paragraph_text = remove_reporter_phrases(next_paragraph_text)
                    next_paragraph_text = remove_inline_figure_table_markers(next_paragraph_text)

                    # ENHANCED: Check for both original format and new metadata format
                    if (is_source_citation(next_paragraph_text) or is_new_metadata_format(next_paragraph_text_original)) and is_valid_headline(text):
                        is_title = True

            if current_section and is_title:
                match_existing_index = re.match(r'^(\d+)\.\s*(.*)', text)
                if match_existing_index:
                    original_title_text, stripped_title_text = text, match_existing_index.group(2).strip()
                    TITLE_MODIFICATIONS.append({'original_text': original_title_text, 'modified_text': stripped_title_text, 'section': current_section, 'original_paragraph_index': i})
                    prospective_title_text = stripped_title_text
                
                section_counters[current_section] = section_counters.get(current_section, 0) + 1
                article_index = section_counters[current_section]
                
                structure['sections'][current_section].append({'index': i, 'text': prospective_title_text, 'type': 'article_title', 'section_index': article_index})
                
                title_cooldown_counter = 1
            else:
                structure['other_content'].append({'index': i, 'text': text, 'type': 'content', 'section': current_section})

    structure['title_format_modifications'] = TITLE_MODIFICATIONS

    with open(json_output_path, 'w', encoding='utf-8') as f:
        json.dump(structure, f, ensure_ascii=False, indent=2, default=str)
    
    return structure

def rebuild_document_from_structure(doc_path, structure_json_path=None, output_path=None, is_monday_mode=False, sunday_date=None):
    """Rebuild document from extracted structure"""

    if structure_json_path is None:
        structure_json_path = doc_path.replace('.docx', '_structure.json')
    if output_path is None:
        output_path = doc_path.replace('.docx', '_final_formatted.docx')
        
    with open(structure_json_path, 'r', encoding='utf-8') as f:
        structure = json.load(f)
        
    new_doc = Document()
    setup_document_fonts(new_doc)

    today_str = datetime.now().strftime("%Y%m%d")
    add_date_line_if_needed(new_doc, today_str)

    editorial_section_header = None
    for content in structure['other_content']:
        if content['type'] == 'section_header' and content['section'] == 'editorial':
            editorial_section_header = content['text']
            break
    if editorial_section_header:
        add_section_header_to_doc(new_doc, editorial_section_header)

    editorial_groups = {g['clean_name']: g for g in structure['editorial_media_groups']}
    for name in EDITORIAL_MEDIA_ORDER:
        if name in editorial_groups:
            if is_monday_mode and sunday_date:
                group = editorial_groups[name]
                # Check if the media group is Sunday news by a stored flag (e.g. `is_sunday_article` from extraction)
                if group.get('is_sunday_article', False):
                    group['clean_name'] = f"{sunday_date} {group['clean_name']}"
            add_media_group_to_document(new_doc, editorial_groups[name])

    all_content = []
    for content in structure['other_content']:
        if content['type'] == 'section_header' and content['section'] == 'editorial':
            continue
        if content['type'] == 'subtitle_removed':
            continue
        all_content.append(('other', content))
    for section_name, articles in structure['sections'].items():
        for article in articles:
            all_content.append(('article', article))
    all_content.sort(key=lambda x: x[1].get('index', x[1].get('start_index', 0)))
    
    previous_was_content = False
    last_article_idx = -1
    st.write("Check: ismondaymode=", is_monday_mode, "sundaydate=", sunday_date)
    for idx, (content_type, content_data) in enumerate(all_content):
        # Debug print
        st.write(f"Processing idx={idx}, type={content_type}, data={content_data}")
        if content_type == 'other' and content_data['type'] == 'section_header':
            section_label = content_data.get('section', '')
            section_text = content_data.get('text', '')
            
            # Debug print
            st.write("Check: ismondaymode=", is_monday_mode, "sundaydate=", sunday_date, "sectionlabel=", section_label)

            # --- Monday Notice Logic: Trigger ONLY before 國際新聞 ---
            if is_monday_mode and sunday_date and (
                section_label == 'international' or '國際新聞' in section_text):
                st.write("[DEBUG] >>> Adding Monday notice before 國際新聞 at idx={idx}")
                add_monday_notice(new_doc, sunday_date)
            add_section_header_to_doc(new_doc, section_text)
            previous_was_content = False
        elif content_type == 'other':
            clean_text = remove_reporter_phrases(content_data['text'])
            clean_text = remove_inline_figure_table_markers(clean_text)
            p = new_doc.add_paragraph(clean_text)
            format_content_paragraph(p)
            previous_was_content = True
        elif content_type == 'article':
            content_data['text'] = remove_reporter_phrases(content_data['text'])
            content_data['text'] = remove_inline_figure_table_markers(content_data['text'])
            add_article_to_document(new_doc, content_data, previous_was_content)
            previous_was_content = True
            last_article_idx = idx


    if last_article_idx != -1:
        add_end_marker(new_doc)

    new_doc.save(output_path)
    return output_path