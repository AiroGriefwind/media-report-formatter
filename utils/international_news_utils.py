
# =============================================================================
# INTERNATIONAL NEWS SPECIFIC FUNCTIONS
# =============================================================================

import re
import time

import tempfile
from datetime import datetime
from docx import Document

from .web_scraping_utils import retry_step, wait_for_search_results, scroll_to_load_all_content, wait_for_ajax_complete
from .document_utils import setup_document_fonts, add_end_marker
from .config import MEDIA_NAME_MAPPINGS

# Web scraping imports
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import (
    TimeoutException,
    ElementClickInterceptedException,
    StaleElementReferenceException,
    WebDriverException,
)

from .wisers_utils import (
    retry_step, 
    wait_for_search_results, 
    scroll_to_load_all_content, 
    wait_for_ajax_complete,
)


@retry_step
def run_international_news_task(**kwargs):
    """Search for international news articles with fallback mechanisms"""
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st = kwargs.get('st_module')
    max_articles = kwargs.get('max_articles', 30)
    # max_words = kwargs.get('max_words', 1000)
    # min_words = kwargs.get('min_words', 200)
    try:
        # Try the saved search approach first
        dropdown_toggle = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "li.dropdown-usersavedquery > a.dropdown-toggle")))
        dropdown_toggle.click()
        time.sleep(3)

        edit_saved_search_btn = wait.until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "button[data-target='#modal-saved-search-ws6']")))
        edit_saved_search_btn.click()
        wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#modal-saved-search-ws6")))
        time.sleep(3)

        # Look for "國際新聞" in the saved searches
        try:
            international_item = wait.until(
                EC.element_to_be_clickable((By.XPATH, "//ul[@class='list-group']//h5[text()='國際新聞']/ancestor::li")))
            international_item.click()
            time.sleep(3)
            
            if st:
                st.write("✅ Found '國際新聞' saved search")
        
        except TimeoutException:
            if st:
                st.warning("⚠️ '國際新聞' saved search not found. Checking available searches...")
            
            # List all available saved searches for debugging
            try:
                search_items = driver.find_elements(By.XPATH, "//ul[@class='list-group']//h5")
                available_searches = [item.text.strip() for item in search_items if item.text.strip()]
                if st and available_searches:
                    st.info(f"Available saved searches: {', '.join(available_searches)}")
            except:
                pass
            
            # Close modal and return empty result
            try:
                close_btn = driver.find_element(By.CSS_SELECTOR, "#modal-saved-search-ws6 .close")
                close_btn.click()
                time.sleep(2)
            except:
                pass
            
            return []

        # Click search button
        search_btn = None
        selectors = [(By.CSS_SELECTOR, "div.modal-footer .btn-default:last-child"),
                    (By.XPATH, "//div[@class='modal-footer']//button[text()='搜索']")]
        
        for selector_type, selector in selectors:
            try:
                search_btn = wait.until(EC.element_to_be_clickable((selector_type, selector)))
                break
            except TimeoutException:
                continue

        if search_btn:
            search_btn.click()
        else:
            driver.execute_script("""
                var buttons = document.querySelectorAll('div.modal-footer button');
                for (var i = 0; i < buttons.length; i++) {
                    if (buttons[i].textContent.trim() === '搜索') {
                        buttons[i].click(); break;
                    }
                }""")

        # Wait for modal to close
        wait.until(EC.invisibility_of_element_located((By.CSS_SELECTOR, "#modal-saved-search-ws6")))
        
        # Wait 15 seconds for search results to fully load
        if st:
            st.write("⏳ Waiting 15 seconds for search results to fully load...")
        time.sleep(15)

        if wait_for_search_results(driver=driver, wait=wait, st_module=st):
            # Scroll to load all content and wait for AJAX
            scroll_to_load_all_content(driver=driver, st_module=st)
            wait_for_ajax_complete(driver, timeout=10)
            
            # Get all article links for scraping
            articles_data = []
            
            for retry in range(3):
                results = driver.find_elements(By.CSS_SELECTOR, 'div.list-group-item.no-excerpt')
                if st:
                    st.write(f"[International News Scrape] Attempt {retry+1}: {len(results)} items found.")
                
                # Limit to around 80-100 articles as requested
                results = results[:max_articles]
                
                for i, result in enumerate(results):
                    try:
                        title_element = result.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a')
                        title = title_element.text.strip()
                        article_url = title_element.get_attribute('href')
                        
                        # Get media name
                        try:
                            media_name_raw = result.find_element(By.CSS_SELECTOR, 'small a').text.strip()
                            mapped_name = next((v for k, v in MEDIA_NAME_MAPPINGS.items() if k in media_name_raw), media_name_raw)
                        except:
                            mapped_name = "Unknown"
                        
                        articles_data.append({
                            'title': title,
                            'url': article_url,
                            'media': mapped_name,
                            'index': i
                        })
                        
                    except Exception as e:
                        if st:
                            st.warning(f"Error extracting article {i}: {e}")
                        continue
                
                if len(articles_data) > 0:
                    break
                time.sleep(2)
            
            return articles_data
        
        return []
        
    except Exception as e:
        if st:
            st.error(f"Error in international news search: {e}")
        return []

def should_scrape_article_based_on_metadata(metadata_text, min_words=200, max_words=1000):
    """Determine if article should be scraped based on metadata - filter out opinion pieces and articles outside word range"""
    if not metadata_text:
        return True  # If no metadata, let it pass to avoid missing articles
    
    # Filter out known opinion/editorial keywords
    opinion_keywords = [
        '社評', '評論', '觀點', '專欄', '分析', '時評', '評論員', '觀察', 
        '社論', '筆記', '隨筆', '札記', '感言', '思考', '反思', '見解',
        'editorial', 'opinion', 'commentary', 'analysis'
    ]
    
    for keyword in opinion_keywords:
        if keyword in metadata_text:
            return False
    
    # Extract word count from metadata
    import re
    word_count_matches = re.findall(r'\|\s*(\d+)\s*字\s*\|', metadata_text)
    if not word_count_matches:
        match = re.search(r'(\d+)\s*字', metadata_text)
        if match:
            word_count = int(match.group(1))
        else:
            return True  # If no word count found, accept by default
    else:
        word_count = int(word_count_matches[0])
    
    # ✅ UPDATED: Filter articles outside the word count range
    if word_count < min_words or word_count > max_words:
        return False
    
    return True

def parse_metadata(raw_meta, title=""):  # 只需 raw_meta，不用 title 做解析
    formatted_meta = raw_meta  # fallback
    
    # 匹配日期 (YYYY-MM-DD 或變體)
    date_match = re.search(r'(\d{4}[-年.]\d{2}[-月.]\d{2})', raw_meta)
    # 匹配字數
    word_match = re.search(r'(\d+)\s*字', raw_meta)
    
    if date_match:
        date_str = date_match.group(1).replace('年', '-').replace('月', '-').replace('日', '').replace('.', '-')
    else:
        date_str = datetime.now().strftime('%Y-%m-%d')
    
    word_str = f"{word_match.group(1)} 字" if word_match else ""
    
    # 提取媒體：剔除日期/字數
    media_part = raw_meta
    if date_match: media_part = media_part.replace(date_match.group(0), '')
    if word_match: media_part = media_part.replace(word_match.group(0), '')
    media_part = media_part.replace('|', '').strip()
    
    # 映射媒體名
    mapped_media = media_part
    if media_part in MEDIA_NAME_MAPPINGS:
        mapped_media = MEDIA_NAME_MAPPINGS[media_part]
    else:
        for k, v in MEDIA_NAME_MAPPINGS.items():
            if k in media_part:
                mapped_media = v
                break
    
    # 組裝：Date | Media | Words
    parts = [p for p in [date_str, mapped_media, word_str] if p]
    return " | ".join(parts) if parts else raw_meta or "No metadata"

def create_hover_preview_report(**kwargs):
    """
    Create a Word document report from hover preview list.
    Expected kwargs:
    - preview_data: list of dict with 'title', 'hover_html', 'hover_text', 'metadata_line'
    - output_path: path to save docx file
    - st_module: optional Streamlit module for logging
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    import re
    from .config import MEDIA_NAME_MAPPINGS 

    preview_data = kwargs.get('preview_data', [])
    output_path = kwargs.get('output_path')
    st = kwargs.get('st_module')

    doc = Document()

    # Add title
    title = doc.add_heading('國際新聞 - 懸停預覽報告', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add date
    doc.add_paragraph(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")\
        .alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"總共找到 {len(preview_data)} 篇文章")
    doc.add_paragraph()

    for item in preview_data:
        # 1. Title (No numbering, just title)
        article_title = item.get('title', 'Unknown')
        doc.add_heading(article_title, level=2)

        # 2. Metadata Line (Date | Media | Words)
        raw_meta = item.get('metadata_line', '')
        formatted_meta = raw_meta  # Default fallback

        # Try to parse and format metadata
        # 尝试匹配日期 (支持 YYYY-MM-DD 或 YYYY.MM.DD 或 YYYY年MM月DD日)
        date_match = re.search(r'(\d{4}[-年.]\d{2}[-月.]\d{2})', raw_meta)
        # 尝试匹配字数
        word_match = re.search(r'(\d+)\s*字', raw_meta)

        # 如果没有日期，默认使用当天 (或者留空，看你需求)
        if date_match:
            date_str = date_match.group(1)
            # 统一日期格式为 YYYY-MM-DD
            date_str = date_str.replace('年', '-').replace('月', '-').replace('日', '').replace('.', '-')
        else:
            # 如果元数据里没日期，可以暂且用当天，或者根据实际情况处理
            # 很多列表页 metadata 只有 "媒体名 字数"
            date_str = datetime.now().strftime('%Y-%m-%d')

        if word_match:
            word_str = f"{word_match.group(1)} 字"
        else:
            word_str = ""

        # 提取媒体名：把日期和字数剔除掉，剩下的就是媒体名
        media_part = raw_meta
        if date_match:
            media_part = media_part.replace(date_match.group(0), '') # 剔除原始日期字符串
        if word_match:
            media_part = media_part.replace(word_match.group(0), '') # 剔除原始字数
        
        # 清理多余符号
        media_part = media_part.replace('|', '').strip()
        
        # 媒体名映射 (Long -> Short)
        mapped_media = media_part
        # Check exact matches first
        if media_part in MEDIA_NAME_MAPPINGS:
            mapped_media = MEDIA_NAME_MAPPINGS[media_part]
        else:
            # Check substring matches
            for k, v in MEDIA_NAME_MAPPINGS.items():
                if k in media_part:
                    mapped_media = v
                    break
        
        # 如果媒体名为空（比如原文只有日期和字数），尝试给一个默认值或者留空
        if not mapped_media and not date_match and not word_match:
             # 极端情况：raw_meta 也是空的
             pass
        elif not mapped_media:
             # 有可能 raw_meta 就是 "明報" 这种纯名字
             # 上面的 replace 逻辑可能会把它处理完
             # 如果剔除后只剩空字符串，说明 raw_meta 本身就是日期或字数？不太可能
             # 这是一个兜底
             pass

        # 重新构建格式: Date | Media | Words
        # 只有当该字段有值时才加入
        parts = []
        if date_str: parts.append(date_str)
        if mapped_media: parts.append(mapped_media)
        if word_str: parts.append(word_str)
        
        if parts:
            formatted_meta = " | ".join(parts)
        
        # Add Metadata Paragraph
        doc.add_paragraph(formatted_meta)
        
        # Optional: Make it look slightly different (e.g. smaller text)
        # meta_para.style = doc.styles['No Spacing'] 

        # 3. Content Body (Cleaned)
        # Get text from hover_text or strip HTML from hover_html
        content_text = item.get('hover_text', '')
        if not content_text:
            import re
            clean_html = re.sub('<[^<]+?>', '\n', item.get('hover_html', ''))
            content_text = clean_html

        # Clean up content: remove Title and Metadata if they appear at the start
        lines = content_text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line: continue
            
            # Remove if line is identical to title
            if line == article_title.strip():
                continue
                
            # Remove if line looks like metadata (contains date and '字')
            if re.search(r'\d{4}[-.]\d{2}[-.]\d{2}', line) and '字' in line:
                continue
                
            cleaned_lines.append(line)
            
        final_content = '\n'.join(cleaned_lines)
        doc.add_paragraph(final_content)

        # Separator
        doc.add_paragraph()

    doc.save(output_path)

    if st:
        st.write(f"✅ Report saved to {output_path}")
    
    return output_path



@retry_step
def scrape_international_articles_sequentially(**kwargs):
    """Scrape international articles with pre-filtering for news only and word count limit"""
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    max_articles = kwargs.get('max_articles', 30)
    max_words = kwargs.get('max_words', 1000)
    st = kwargs.get('st_module')

    scraped_articles = []
    skipped_articles = []
    original_window = driver.current_window_handle

    # Get all article elements from current search results page
    results = driver.find_elements(By.CSS_SELECTOR, 'div.list-group-item.no-excerpt')
    results = results[:max_articles]  # Limit to max articles

    if st:
        st.write(f"Found {len(results)} articles to filter and scrape sequentially")

    scraped_count = 0
    for idx in range(len(results)):
        try:
            # IMPORTANT: Re-find elements each iteration because DOM may change
            current_results = driver.find_elements(By.CSS_SELECTOR, 'div.list-group-item.no-excerpt')
            if idx >= len(current_results):
                break

            result = current_results[idx]

            # ✅ NEW: Pre-filter using metadata from search results page
            try:
                metadata_preview = result.find_element(By.CSS_SELECTOR, 'small').text.strip()
                title_element = result.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a')
                title_preview = title_element.text.strip()
            except:
                metadata_preview = None
                title_preview = f"Article {idx+1}"

            # Apply filter before scraping
            if metadata_preview and not should_scrape_article_based_on_metadata(metadata_preview, max_words=max_words):
                if st:
                    # Extract word count for logging
                    word_match = re.search(r'(\d+)\s*字', metadata_preview)
                    word_count = word_match.group(1) if word_match else "unknown"
                    st.write(f"⏭️ Skipping article {idx+1}/{len(results)} (Word count: {word_count}, Title: {title_preview[:30]}...)")
                
                skipped_articles.append({
                    'title': title_preview,
                    'metadata': metadata_preview,
                    'reason': 'Opinion piece or over word limit'
                })
                continue

            if st:
                st.write(f"✅ Scraping article {scraped_count+1} (Index {idx+1}/{len(results)}): {title_preview[:50]}...")

            # Click the article link (same as author search)
            article_link = result.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a')
            article_link.click()

            # Wait for new window and switch to it (same as author search)
            wait.until(EC.number_of_windows_to_be(2))
            for window_handle in driver.window_handles:
                if window_handle != original_window:
                    driver.switch_to.window(window_handle)
                    break

            # Scrape content from article detail page
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div.article-detail')))
            time.sleep(2)

            # Extract title from article page
            title = driver.find_element(By.CSS_SELECTOR, 'h3').text.strip()

            # Extract full metadata from article page
            try:
                full_metadata_element = driver.find_element(By.CSS_SELECTOR, 'div.article-subheading')
                full_metadata_line = full_metadata_element.text.strip()
            except:
                full_metadata_line = metadata_preview or "Unknown Source"

            # Extract article content
            paragraphs = []
            content_elements = driver.find_elements(By.CSS_SELECTOR, 'div.description p')
            for p in content_elements:
                text = p.text.strip()
                if text:
                    paragraphs.append(text)

            content_body = '\n\n'.join(paragraphs) if paragraphs else ""

            article_data = {
                'title': title,
                'metadata_line': full_metadata_line,
                'content': content_body,
                'full_text': f"{title}\n\n{full_metadata_line}\n\n{content_body}" if content_body else f"{title}\n\n{full_metadata_line}"
            }

            scraped_articles.append(article_data)
            scraped_count += 1

            # Close current tab and return to original window (same as author search)
            driver.close()
            driver.switch_to.window(original_window)

            time.sleep(1)  # Brief pause between articles

        except Exception as e:
            if st:
                st.warning(f"Failed to scrape article {idx+1}: {e}")

            # Ensure we're back on the original window
            try:
                driver.switch_to.window(original_window)
            except:
                pass
            continue

    # Log filtering results
    if st:
        st.write(f"📊 Filtering Summary:")
        st.write(f"• Total articles found: {len(results)}")
        st.write(f"• Articles scraped: {len(scraped_articles)}")
        st.write(f"• Articles skipped: {len(skipped_articles)}")

    return scraped_articles, skipped_articles

def scrape_specific_articles_by_indices(driver, wait, articles_to_scrape, st_module=None):
    """
    Scrape full content for a specific list of articles based on their original index on the search page.
    
    articles_to_scrape: List of dicts, must contain 'original_index' (int) and 'title' (str).
    """
    scraped_data = []
    original_window = driver.current_window_handle
    
    total = len(articles_to_scrape)
    
    if st_module:
        progress_bar = st_module.progress(0)
        status_text = st_module.empty()
    
    for i, item in enumerate(articles_to_scrape):
        idx = item.get('original_index')
        title = item.get('title', 'Unknown')
        
        if idx is None:
            if st_module: st_module.warning(f"Skipping {title}: No original index found.")
            continue
            
        try:
            if st_module:
                status_text.text(f"Scraping {i+1}/{total}: {title}...")
                progress_bar.progress((i) / total)
            
            # 1. Re-locate the element on the search results page
            # We need to find all items again to ensure freshness, then pick the [idx] one
            all_results = driver.find_elements(By.CSS_SELECTOR, 'div.list-group-item.no-excerpt')
            
            if idx >= len(all_results):
                if st_module: st_module.warning(f"Index {idx} is out of range (Total: {len(all_results)}). Skipping.")
                continue
                
            target_element = all_results[idx]
            
            # 2. Click logic (Same as scrape_international_articles_sequentially)
            article_link = target_element.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a')
            article_link.click()
            
            # 3. Handle window switching
            wait.until(EC.number_of_windows_to_be(2))
            for window_handle in driver.window_handles:
                if window_handle != original_window:
                    driver.switch_to.window(window_handle)
                    break
            
            # 4. Scrape content
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div.article-detail')))
            time.sleep(1.5)
            
            # Extract Data
            full_title = driver.find_element(By.CSS_SELECTOR, 'h3').text.strip()
            try:
                meta_el = driver.find_element(By.CSS_SELECTOR, 'div.article-subheading')
                full_metadata = meta_el.text.strip()
            except:
                full_metadata = ""
                
            paragraphs = [p.text.strip() for p in driver.find_elements(By.CSS_SELECTOR, 'div.description p') if p.text.strip()]
            content_body = '\n\n'.join(paragraphs)
            
            article_data = {
                'title': full_title,
                'metadata_line': full_metadata,
                'content': content_body,
                'full_text': f"{full_title}\n\n{full_metadata}\n\n{content_body}"
            }
            
            # Preserve AI analysis if it exists
            if 'ai_analysis' in item:
                article_data['ai_analysis'] = item['ai_analysis']
                
            scraped_data.append(article_data)
            
            # 5. Close and return
            driver.close()
            driver.switch_to.window(original_window)
            time.sleep(0.5)
            
        except Exception as e:
            if st_module: st_module.error(f"Failed to scrape {title}: {e}")
            try:
                driver.switch_to.window(original_window)
            except:
                pass
    
    if st_module:
        progress_bar.progress(1.0)
        status_text.text("Scraping complete!")
        
    return scraped_data

def extract_news_id_from_html(hover_html):
    """Extract news_id from hover_html content"""
    import re
    if not hover_html:
        return None
    
    # Match patterns like: id="news:2500^202512179501277(S:568259853)"
    match = re.search(r'id="(news:[^"]+)"', hover_html)
    if match:
        return match.group(1)
    return None


def _safe_click(driver, element, st_module=None, attempts: int = 3) -> None:
    """
    Click an element robustly on pages with fixed headers/overlays.
    - Scroll element into viewport center to avoid header covering it
    - Retry on click interception / transient webdriver issues
    - Fallback to JS click
    """
    last_err = None
    for n in range(attempts):
        try:
            driver.execute_script(
                "arguments[0].scrollIntoView({block:'center', inline:'nearest'});",
                element,
            )
            # Offset a bit in case a fixed header still overlaps
            driver.execute_script("window.scrollBy(0, -120);")
            time.sleep(0.15)

            try:
                ActionChains(driver).move_to_element(element).pause(0.05).perform()
            except Exception:
                # Not critical; proceed to click
                pass

            element.click()
            return
        except (ElementClickInterceptedException, WebDriverException) as e:
            last_err = e
            try:
                driver.execute_script(
                    "arguments[0].scrollIntoView({block:'center', inline:'nearest'});",
                    element,
                )
                driver.execute_script("window.scrollBy(0, -120);")
                time.sleep(0.15)
                driver.execute_script("arguments[0].click();", element)
                return
            except Exception as e2:
                last_err = e2
                time.sleep(0.2)
        except StaleElementReferenceException as e:
            # Caller should re-locate the element
            last_err = e
            break

    if st_module and last_err:
        st_module.write(f"⚠️ 点击重试仍失败: {last_err}")
    raise last_err if last_err else RuntimeError("safe_click failed without exception")


def scrape_articles_by_news_id(driver, wait, articles_to_scrape, st_module=None):
    """
    Scrape articles by locating them via news_id (preferred) or title (fallback).
    This avoids index-based issues when search results change.
    
    articles_to_scrape: List of dicts with 'news_id', 'title', 'formatted_metadata'
    """
    scraped_data = []
    original_window = driver.current_window_handle
    total = len(articles_to_scrape)
    
    if st_module:
        progress_bar = st_module.progress(0)
        status_text = st_module.empty()
    
    for i, item in enumerate(articles_to_scrape):
        news_id = item.get('news_id')
        title = item.get('title', 'Unknown')
        metadata = item.get('formatted_metadata', '')
        multi_newspapers = bool(item.get("multi_newspapers", False))
        
        try:
            if st_module:
                status_text.text(f"🔍 定位 {i+1}/{total}: {title[:40]}...")
                progress_bar.progress(i / total)
            
            # Strategy 1: Try to locate by news_id
            target_element = None
            
            if news_id:
                try:
                    # Find element with this news_id in the search results page
                    # The news_id is usually in the hover content div
                    news_id_element = driver.find_element(By.XPATH, f"//*[@id='{news_id}']")
                    
                    # Navigate up to the parent list item
                    target_element = news_id_element.find_element(By.XPATH, 
                        "./ancestor::div[contains(@class, 'list-group-item')]")
                    
                    if st_module:
                        st_module.write(f"✅ 用 news_id 定位成功: {title[:40]}")
                except Exception as e:
                    if st_module:
                        st_module.write(f"⚠️ news_id 定位失败，尝试标题匹配: {title[:40]}")
            
            # Strategy 2: Fallback to title matching
            if not target_element:
                try:
                    # Normalize spaces and try exact match first
                    normalized_title = ' '.join(title.split())
                    
                    # Try exact match
                    xpath_exact = f"//h4[@class='list-group-item-heading']//a[normalize-space()='{normalized_title}']"
                    title_links = driver.find_elements(By.XPATH, xpath_exact)
                    
                    if not title_links:
                        # Try contains match (for truncated titles)
                        xpath_contains = f"//h4[@class='list-group-item-heading']//a[contains(normalize-space(), '{normalized_title[:30]}')]"
                        title_links = driver.find_elements(By.XPATH, xpath_contains)
                    
                    if title_links:
                        # Get the parent list item
                        target_element = title_links[0].find_element(By.XPATH, 
                            "./ancestor::div[contains(@class, 'list-group-item')]")
                        
                        if st_module:
                            st_module.write(f"✅ 用标题定位成功: {title[:40]}")
                    else:
                        if st_module:
                            st_module.warning(f"❌ 无法定位文章: {title}")
                        continue
                        
                except Exception as e:
                    if st_module:
                        st_module.error(f"❌ 定位失败: {title} - {e}")
                    continue
            
            # Now click the article link (same logic as before)
            if target_element:
                article_link = target_element.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a')
                # Robust click: overlays like <div class="subtitle"> can intercept clicks
                handles_before = list(driver.window_handles)
                _safe_click(driver, article_link, st_module=st_module, attempts=3)
                
                # Wait for new window and switch
                WebDriverWait(driver, 8).until(
                    lambda d: len(d.window_handles) == len(handles_before) + 1
                )
                for window_handle in driver.window_handles:
                    if window_handle != original_window:
                        driver.switch_to.window(window_handle)
                        break
                
                # Scrape content
                wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div.article-detail')))
                time.sleep(1.5)
                
                # Extract data
                full_title = driver.find_element(By.CSS_SELECTOR, 'h3').text.strip()
                
                try:
                    meta_el = driver.find_element(By.CSS_SELECTOR, 'div.article-subheading')
                    full_metadata = meta_el.text.strip()
                except:
                    full_metadata = metadata
                
                paragraphs = [p.text.strip() for p in driver.find_elements(By.CSS_SELECTOR, 'div.description p') 
                             if p.text.strip()]
                content_body = '\n\n'.join(paragraphs)
                
                article_data = {
                    'title': full_title,
                    'metadata_line': full_metadata,
                    'content': content_body,
                    'full_text': f"{full_title}\n\n{full_metadata}\n\n{content_body}",
                    # carry UI flags forward
                    'multi_newspapers': multi_newspapers,
                }
                
                # Preserve AI analysis if exists
                if 'ai_analysis' in item:
                    article_data['ai_analysis'] = item['ai_analysis']
                
                scraped_data.append(article_data)
                
                # Close and return to search results
                driver.close()
                driver.switch_to.window(original_window)
                time.sleep(0.5)
                
        except Exception as e:
            if st_module:
                st_module.error(f"❌ 爬取失败: {title[:40]} - {e}")
            try:
                driver.switch_to.window(original_window)
            except:
                pass
    
    if st_module:
        progress_bar.progress(1.0)
        status_text.text(f"✅ 爬取完成! 成功: {len(scraped_data)}/{total}")
    
    return scraped_data


def _inject_multi_newspaper_placeholder(metadata_line: str) -> str:
    """
    Inject '==' into the first part of a metadata line so that
    utils.document_utils.transform_metadata_line() will render '及多份報章'.
    We intentionally keep this lightweight and non-destructive.
    """
    if not metadata_line:
        return metadata_line
    if "==" in metadata_line:
        return metadata_line

    # Preferred: metadata lines with pipes where the first part is "media page section ..."
    if "|" in metadata_line:
        parts = metadata_line.split("|")
        parts[0] = parts[0].rstrip() + " =="
        return "|".join(parts)

    # Fallback: no pipes, append token at end.
    return metadata_line.rstrip() + " =="


@retry_step
def create_international_news_report(**kwargs):
    """Create Word document report for international news"""
    articles_data = kwargs.get('articles_data')
    output_path = kwargs.get('output_path')
    st = kwargs.get('st_module')

    # 防禦：如果沒有有效 articles_data，提早拋錯，避免 NoneType 迭代
    if not articles_data:
        raise ValueError("create_international_news_report: 'articles_data' is empty or None")

    from docx import Document

    doc = Document()
    setup_document_fonts(doc)

    # Add title
    doc.add_heading('國際新聞摘要', level=1)
    doc.add_paragraph()

    # Add date
    today_str = datetime.now().strftime("%Y年%m月%d日")
    date_para = doc.add_paragraph(f"日期：{today_str}")
    date_para.add_run().add_break()

    # Add articles
    for i, article in enumerate(articles_data, 1):
        if article and article.get('full_text'):
            # Add article number and title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{i}. {article['title']}")
            title_run.bold = True

            # ✅ CHANGED: Add full metadata line from article page
            if article.get('metadata_line'):
                metadata_line = article['metadata_line']
                if bool(article.get("multi_newspapers", False)):
                    metadata_line = _inject_multi_newspaper_placeholder(metadata_line)
                metadata_para = doc.add_paragraph(metadata_line)
                metadata_para.style = doc.styles['Normal']

            # Add content
            if article.get('content'):
                for paragraph_text in article['content'].split('\n\n'):
                    if paragraph_text.strip():
                        doc.add_paragraph(paragraph_text.strip())

            # Add spacing between articles
            doc.add_paragraph()

    # Add end marker
    add_end_marker(doc)

    doc.save(output_path)
    return output_path

