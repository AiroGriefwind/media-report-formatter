import streamlit as st
import tempfile
import os
import shutil
import traceback
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from io import BytesIO
import time
import base64
import requests

#Import for retrying functions
from functools import wraps

# Import for login error handling
from selenium.common.exceptions import NoSuchElementException

import traceback
from collections import defaultdict
import re
import json

# Document processing imports
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from opencc import OpenCC
import glob
from pathlib import Path

# Web scraping imports
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import TimeoutException
from twocaptcha import TwoCaptcha

# =============================================================================
# CONFIGURATION CONSTANTS
# =============================================================================

# Add specific incorrect-to-correct mappings here.
# This runs AFTER the main s2hk conversion.
CORRECTION_MAP = {
    "餘錦賢": "余錦賢",
    # Add other corrections here, e.g., "错误词": "正确词"
}

# Editorial media order
EDITORIAL_MEDIA_ORDER = [
    '商報', '大公', '文匯', '東方', '星島', '信報', '明報', '經濟', '成報', '頭條', 'am730', 'SCMP'
]

# Universal media name mappings
MEDIA_NAME_MAPPINGS = {
    '信報財經新聞': '信報', '信報': '信報', '明報': '明報', '頭條日報': '頭條', '文匯報': '文匯', '成報': '成報',
    '香港經濟日報': '經濟', '經濟日報': '經濟', '東方日報': '東方', '香港商報': '商報', '商報': '商報', '大公報': '大公',
    '星島日報': '星島', 'Am730': 'am730', '南華早報': 'SCMP', 'SCMP': 'SCMP'
}

# Editorial media names
EDITORIAL_MEDIA_NAMES = [
    '信報', '明報', '頭條', '文匯', '成報', '經濟', '東方', '商報', '大公', '星島', 'am730', 'SCMP'
]

# Global list for title modifications
TITLE_MODIFICATIONS = []

# Web scraping URL
WISERS_URL = 'https://login.wisers.net/'

# =============================================================================
# RETRY DECORATOR
# =============================================================================


def retry_step(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        st = kwargs.get('st_module')
        driver = kwargs.get('driver')
        retry_limit = 3
        for trial in range(1, retry_limit + 1):
            try:
                result = func(*args, **kwargs)
                if st:
                    st.write(f"✅ Step {func.__name__} succeeded on attempt {trial}")
                return result
            except Exception as e:
                if st:
                    st.warning(f"⚠️ Step {func.__name__} failed on attempt {trial}: {e}")
                if driver and st:
                    try:
                        img_bytes = driver.get_screenshot_as_png()
                        st.image(img_bytes, caption=f"Screencap after failure in {func.__name__}, attempt {trial}")
                        st.download_button(
                            label=f"Download {func.__name__}_attempt{trial}_screenshot.png",
                            data=img_bytes,
                            file_name=f"{func.__name__}_attempt{trial}_screenshot.png",
                            mime="image/png"
                        )
                    except Exception as screencap_err:
                        if st:
                            st.warning(f"Screencap failed: {screencap_err}")
                time.sleep(2)
                if trial == retry_limit:
                    if st:
                        st.error(f"❌ Step {func.__name__} failed after {retry_limit} attempts.")
                    # Perform robust forced logout after maximum retries
                    try:
                        if driver:
                            robust_logout_request(driver=driver, st_module=st)
                        elif st:
                            st.warning("Driver not available for robust logout request.")
                    except Exception as logout_err:
                        if st:
                            st.warning(f"Robust logout request failed: {logout_err}")
                    raise Exception(f"Step {func.__name__} failed after {retry_limit} attempts.")
    return wrapper

# =============================================================================
# DOCUMENT FORMATTING FUNCTIONS
# =============================================================================

def is_source_citation(text):
    """Check if text is a source citation"""
    if not text: 
        return False
    if ']' in text and text.index(']') < 30: 
        return False
    if re.match(r'^.{1,20}[:：]', text): 
        return True
    common_media_prefixes = "|".join(re.escape(k) for k in MEDIA_NAME_MAPPINGS.keys())
    if re.match(rf'^({common_media_prefixes})\s*[:：]', text): 
        return True
    return False

def is_valid_headline(text):
    """Validates if a line of text could be a headline"""
    if not text or len(text.strip()) < 5:
        return False
    if re.search(r'[，,。]', text):
        return False
    if re.search(r'[.?!]$', text.strip()):
        return False
    if ']' in text:
        return False
    return True

def is_new_metadata_format(text):
    """
    Detects if a line is in the new metadata format.
    Format: "media name + page number + section name | word count | date"
    """
    if not text:
        return False
    
    # Check for exactly two '|' characters
    if text.count('|') != 2:
        return False
    
    # Check that it doesn't end with punctuation
    if re.search(r'[。，.，]$', text.strip()):
        return False
    
    # Additional validation - should have format like "媒体名 页码 栏目名 |字数 |日期"
    parts = text.split('|')
    if len(parts) != 3:
        return False
    
    # The first part should contain media name and page info
    first_part = parts[0].strip()
    if not first_part:
        return False
    
    # The second part should contain word count (should have '字')
    second_part = parts[1].strip()
    if '字' not in second_part:
        return False
    
    # The third part should look like a date
    third_part = parts[2].strip()
    if not re.match(r'^\d{4}-\d{2}-\d{2}$', third_part):
        return False
    
    return True

def transform_metadata_line(metadata_text, next_paragraph_text):
    """
    Transforms metadata line from new format to semicolon format.
    From: "香港经济日报 A04 评析天下 |911 字 |2025-07-16"
    To: "经济 A04：全国政协副主席梁振英出任主席的共享基金会，主力为一带一路国家提供"
    """
    if not metadata_text or not next_paragraph_text:
        return metadata_text
    
    # Split by '|' to get the main part
    parts = metadata_text.split('|')
    if len(parts) < 1:
        return metadata_text
    
    next_paragraph_text = remove_reporter_phrases(next_paragraph_text)
    main_part = parts[0].strip()
    
    # Detect full media name, page, and placeholder '=='
    #    Regex: (media) (page)(==)? (optional section)
    m = re.match(
        r'^([\u4e00-\u9fa5A-Za-z（）()]+)\s+([A-Z]\d{2})(==)?(?:\s+[^\s]+)?',
        main_part
    )
    if not m:
        # Fallback to original if it doesn’t match
        return metadata_text
    
    media_name = m.group(1)
    page_number = m.group(2)
    has_placeholder  = bool(m.group(3))
    
    # Convert long media name to short name using enhanced mapping
    short_media_name = get_short_media_name(media_name)

    #  Rebuild the “及多份報章” phrase if the placeholder was present
    suffix = '及多份報章' if has_placeholder else ''
    
    # Extract first paragraph
    body = next_paragraph_text.strip()

    transformed = f"{short_media_name} {page_number}{suffix}：{body}"
    return transformed
    

def get_short_media_name(full_media_name):
    """
    Convert full media name to short name with flexible matching.
    """
    # Direct mapping first
    if full_media_name in MEDIA_NAME_MAPPINGS:
        return MEDIA_NAME_MAPPINGS[full_media_name]
    
    # Flexible matching for common patterns (handles both simplified and traditional)
    if '经济日报' in full_media_name or '經濟日報' in full_media_name:
        return '经济'
    if '明报' in full_media_name or '明報' in full_media_name:
        return '明报'
    if '文汇报' in full_media_name or '文匯報' in full_media_name:
        return '文汇'
    if '东方日报' in full_media_name or '東方日報' in full_media_name:
        return '东方'
    if '星岛日报' in full_media_name or '星島日報' in full_media_name:
        return '星岛'
    if '大公报' in full_media_name or '大公報' in full_media_name:
        return '大公'
    if '头条' in full_media_name or '頭條' in full_media_name:
        return '头条'
    if '成报' in full_media_name or '成報' in full_media_name:
        return '成报'
    if '商报' in full_media_name or '商報' in full_media_name:
        return '商报'
    if 'am730' in full_media_name.lower():
        return 'am730'
    if '南華早報' in full_media_name or '南华早报' in full_media_name:
        return 'SCMP'
    
    # If no match found, return the original name
    return full_media_name

def remove_reporter_phrases(text):
    if not text:
        return ""
    #Remove `●香港文匯報記者 or 香港文汇报记者` and anything after it
    text = re.sub(r'(●香港文匯報記者|●香港文汇报记者).*$', '', text, flags=re.MULTILINE)

    #Remove reporting agency content between first colon and '报道：' or '報道：'
    match = re.search(r'(^.+?：)', text)
    if match:
        prefix = match.group(1)  # everything up to and including the first colon
        rest = text[match.end():]  # everything after the first colon
        
        # Find 报道： or 報道： in the rest
        rep_match = re.search(r'(报道：|報道：)', rest)
        if rep_match:
            # Remove text from start of rest up to and including 报道：/報道：
            rest = rest[rep_match.end():].lstrip()
            text = prefix + rest

    # Remove 【...】 containing keywords
    pattern_brackets = r'【[^】]*?(记者|記者|报道|報道|报讯|報訊)[^】]*?】'
    text = re.sub(pattern_brackets, '', text)
    # Remove （...） containing keywords, using a function for precision
    def paren_replacer(match):
        if re.search(r'(记者|記者|报道|報道|报讯|報訊)', match.group(1)):
            return ''
        return match.group(0)
    text = re.sub(r'（([^）]*)）', paren_replacer, text)
    # Remove the fixed phrase
    text = text.replace('香港文汇报訊', '').replace('香港文匯報訊', '')
    return text.strip()

def convert_to_traditional_chinese(text):
    """Convert simplified Chinese to traditional Chinese"""
    if not text or not text.strip():
        return text
    try:
        cc = OpenCC('s2hk')
        converted_text = cc.convert(text)
        return converted_text
    except Exception as e:
        print(f"Warning: Chinese conversion failed for text: {text[:50]}... Error: {str(e)}")
        return text
    
def apply_gatekeeper_corrections(text):
    """
    Applies a second round of specific corrections based on the CORRECTION_MAP.
    """
    if not text:
        return text
    for error, correction in CORRECTION_MAP.items():
        if error in text:
            text = text.replace(error, correction)
    return text

def setup_document_fonts(doc):
    """Setup document fonts"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def add_first_page_header(doc, logo_path):
    """Add header only on the first page"""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    header = section.first_page_header
    header_para = header.paragraphs[0]
    header_para.clear()
    
    left_run = header_para.add_run("亞聯每日報章摘要")
    left_run.font.name = 'Calibri'
    left_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    left_run.font.size = Pt(18)
    
    tab_stops = header_para.paragraph_format.tab_stops
    tab_stops.clear_all()
    tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    
    header_para.add_run("\t\t")
    logo_run = header_para.add_run()
    logo_run.add_picture(logo_path, width=Cm(5.95), height=Cm(2.04))
    header_para.style = doc.styles['Header']

def add_first_page_footer(doc):
    """Add footer only on the first page"""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    footer = section.first_page_footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    
    footer_lines = [
        "香港金鐘夏愨道18號海富中心24樓  電話: 2114 4960  傳真: 3544 2933",
        "電郵: info@asianet-sprg.com.hk", 
        "網頁: http://www.asianet-sprg.com.hk"
    ]
    
    for i, line in enumerate(footer_lines):
        run = footer_para.add_run(line)
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(12)
        if i < len(footer_lines) - 1:
            footer_para.add_run('\n')
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_subsequent_pages_header(doc):
    """Add header to pages 2 onwards"""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.clear()
    
    header_run = header_para.add_run("AsiaNet亞聯政經顧問")
    header_run.font.name = 'Calibri'
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    header_run.font.size = Pt(12)
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_subsequent_pages_footer(doc):
    """Add footer to pages 2 onwards with page numbers"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    
    footer_run = footer_para.add_run()
    footer_run.font.name = 'Calibri'
    footer_run.font.size = Pt(12)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    footer_run._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    footer_run._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    footer_run._element.append(fldChar2)
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def detect_section_type(text):
    """Detect the type of section from text"""
    if not text: 
        return None
    sections = {
        'editorial': r'^報章社評\s*$', 
        'international': r'^國際新聞[:：]?\s*$', 
        'china': r'^大中華新聞\s*$', 
        'local': r'^本地新聞\s*$', 
        'financial': r'^財經新聞\s*$', 
        'Hong Kong': r'^香港本地新聞\s*$', 
        'entertainment': r'^娛樂新聞\s*$', 
        'sports': r'^體育新聞\s*$', 
        'property': r'^地產新聞\s*$'
    }
    for name, pattern in sections.items():
        if re.match(pattern, text): 
            return name
    return None

def detect_editorial_media_line(text):
    """Detects editorial media lines"""
    if not text: 
        return None
    
    match = re.match(r'^([^：]+)：(.*)$', text)
    if match:
        potential_name, content = match.group(1).strip(), match.group(2).strip()
        
        if potential_name in EDITORIAL_MEDIA_NAMES:
            return {'full_name': potential_name, 'clean_name': potential_name, 'content': content}
        
        if potential_name in MEDIA_NAME_MAPPINGS:
            clean_name = MEDIA_NAME_MAPPINGS[potential_name]
            return {'full_name': potential_name, 'clean_name': clean_name, 'content': content}
    
    return None

def is_editorial_continuation(text):
    """Detects if a line is a continuation of editorial content"""
    if not text: 
        return False
    if re.match(r'^\s*\d+\.\s+', text): 
        return True
    if re.match(r'^[\t\s]{2,}', text): 
        return True
    if len(text.strip()) > 15: 
        return True
    return False

def format_content_paragraph(paragraph):
    """Format content paragraph"""
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

def format_media_first_line_hanging(paragraph, label_length):
    """Format media first line with hanging indent"""
    pf = paragraph.paragraph_format
    indent_amount = Pt(54)
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.keep_with_next = True
    pf.left_indent = indent_amount
    pf.first_line_indent = -indent_amount
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(indent_amount)

def format_section_header(paragraph):
    """Format section header"""
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.left_indent = Pt(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True

def format_article_title(paragraph, needs_spacing):
    """Format article title"""
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(12) if needs_spacing else Pt(0)
    pf.space_after = Pt(0)
    pf.keep_with_next = True

def add_section_header_to_doc(doc, text):
    """Add section header to document"""
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    format_section_header(p)

def add_article_to_document(doc, article_data, needs_spacing):
    """Add article to document"""
    p = doc.add_paragraph(article_data['text'], style='List Number')
    p.runs[0].bold = True
    #p.style = doc.styles['Normal']
    format_article_title(p, needs_spacing)

def add_media_group_to_document(new_doc, media_group):
    """Add media group to document"""
    media_label = f"{media_group['clean_name']}："
    label_length = len(media_label)
    full_width_space = '\u3000'

    para = new_doc.add_paragraph()
    para.add_run(f"{media_label}{media_group['first_item']}")
    format_media_first_line_hanging(para, label_length)

    for item in media_group['additional_items']:
        item_para = new_doc.add_paragraph()
        item_para.add_run(full_width_space * label_length + item['text'])
        format_media_first_line_hanging(item_para, label_length)

def add_date_line_if_needed(doc, date_str):
    """
    Add date line if needed, safely handling empty documents.
    This corrected version avoids the index error.
    """
    # First, check if the document already has a date line.
    if doc.paragraphs:
        first_text_paragraph = next((p for p in doc.paragraphs if p.text.strip()), None)
        if first_text_paragraph and re.match(r'^\d{8}$', first_text_paragraph.text.strip()):
            return  # Date already exists, so we do nothing.

    # If the document is empty or the date is missing, we add it.
    # This method safely adds the paragraph and moves it to the beginning.
    p = doc.add_paragraph(date_str)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_element = p._p
    doc._body._element.remove(p_element)
    doc._body._element.insert(0, p_element)


def add_end_marker(doc):
    """Add end marker to document"""
    blank_para = doc.add_paragraph("")
    blank_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    end_para = doc.add_paragraph("（完）")
    end_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    end_para.style = doc.styles['Normal']

def extract_document_structure(doc_path, json_output_path=None):
    """
    Extracts structure using state-based logic with Chinese conversion.
    """
    global TITLE_MODIFICATIONS
    TITLE_MODIFICATIONS = []

    if json_output_path is None:
        json_output_path = doc_path.replace('.docx', '_structure.json')
    
    doc = Document(doc_path)
    structure = {'total_paragraphs': len(doc.paragraphs), 'editorial_media_groups': [], 'sections': {}, 'other_content': []}
    
    current_section = None
    in_editorial = False
    section_counters = {}
    
    is_expecting_title = False
    title_cooldown_counter = 0

    paragraphs = doc.paragraphs
    num_paragraphs = len(paragraphs)
    
    # add paragraph while skipping first paragraph after metadata
    skip_next = False
    for i, paragraph in enumerate(paragraphs):
        if skip_next:
            skip_next = False
            continue  # Skip the paragraph after the metadata line

        
        original_text = paragraph.text.strip()
        text = convert_to_traditional_chinese(original_text)
        text = apply_gatekeeper_corrections(text)

        text = remove_reporter_phrases(text)

        if is_new_metadata_format(original_text):
            next_content = ""
            if i + 1 < num_paragraphs:
                next_paragraph_text = paragraphs[i + 1].text.strip()
                next_content = convert_to_traditional_chinese(next_paragraph_text)
                next_content = apply_gatekeeper_corrections(next_content)
            # Transform the metadata line with the *lead* of the next paragraph
            text = transform_metadata_line(text, next_content)
            skip_next = True   # Set flag to skip the next paragraph

        section_type = detect_section_type(text)
        if section_type:
            current_section = section_type
            in_editorial = (section_type == 'editorial')
            is_expecting_title = not in_editorial
            title_cooldown_counter = 0
            if section_type not in structure['sections']:
                structure['sections'][section_type] = []
            structure['other_content'].append({'index': i, 'text': text, 'type': 'section_header', 'section': section_type})
            continue
        
        if not text: continue
        
        if in_editorial:
            media_info = detect_editorial_media_line(text)
            if media_info:
                # Apply conversion to media content as well
                converted_content = convert_to_traditional_chinese(media_info['content'])
                # Apply gatekeeper corrections to media content
                converted_content = apply_gatekeeper_corrections(converted_content)
                # Remove reporter phrases from media content
                converted_content = remove_reporter_phrases(converted_content)

                media_info['content'] = converted_content
                current_media_group = {'clean_name': media_info['clean_name'], 'original_name': media_info['full_name'], 'start_index': i, 'first_item': converted_content, 'additional_items': []}
                structure['editorial_media_groups'].append(current_media_group)
            elif 'current_media_group' in locals() and current_media_group and is_editorial_continuation(text):
                current_media_group['additional_items'].append({'index': i, 'text': text})
        else:
            # Rest of the existing logic remains the same, but with converted text
            if title_cooldown_counter > 0:
                structure['other_content'].append({'index': i, 'text': text, 'type': 'content', 'section': current_section})
                title_cooldown_counter -= 1
                continue

            is_title = False
            prospective_title_text = text

            if is_expecting_title:
                is_title = True
                is_expecting_title = False
            else:
                if i + 1 < num_paragraphs:
                    next_paragraph_text_original = paragraphs[i+1].text.strip()
                    next_paragraph_text = convert_to_traditional_chinese(next_paragraph_text_original)
                    next_paragraph_text = apply_gatekeeper_corrections(next_paragraph_text)
                    
                    # ENHANCED: Check for both original format and new metadata format
                    if (is_source_citation(next_paragraph_text) or is_new_metadata_format(next_paragraph_text_original)) and is_valid_headline(text):
                        is_title = True

            if current_section and is_title:
                match_existing_index = re.match(r'^(\d+)\.\s*(.*)', text)
                if match_existing_index:
                    original_title_text, stripped_title_text = text, match_existing_index.group(2).strip()
                    TITLE_MODIFICATIONS.append({'original_text': original_title_text, 'modified_text': stripped_title_text, 'section': current_section, 'original_paragraph_index': i})
                    prospective_title_text = stripped_title_text
                
                section_counters[current_section] = section_counters.get(current_section, 0) + 1
                article_index = section_counters[current_section]
                
                structure['sections'][current_section].append({'index': i, 'text': prospective_title_text, 'type': 'article_title', 'section_index': article_index})
                
                title_cooldown_counter = 1
            else:
                structure['other_content'].append({'index': i, 'text': text, 'type': 'content', 'section': current_section})

    structure['title_format_modifications'] = TITLE_MODIFICATIONS

    with open(json_output_path, 'w', encoding='utf-8') as f:
        json.dump(structure, f, ensure_ascii=False, indent=2, default=str)
    
    return structure

def rebuild_document_from_structure(doc_path, structure_json_path=None, output_path=None):
    """Rebuild document from extracted structure"""
    if structure_json_path is None:
        structure_json_path = doc_path.replace('.docx', '_structure.json')
    if output_path is None:
        output_path = doc_path.replace('.docx', '_final_formatted.docx')
        
    with open(structure_json_path, 'r', encoding='utf-8') as f:
        structure = json.load(f)
        
    new_doc = Document()
    setup_document_fonts(new_doc)

    today_str = datetime.now().strftime("%Y%m%d")
    add_date_line_if_needed(new_doc, today_str)

    editorial_section_header = None
    for content in structure['other_content']:
        if content['type'] == 'section_header' and content['section'] == 'editorial':
            editorial_section_header = content['text']
            break
    if editorial_section_header:
        add_section_header_to_doc(new_doc, editorial_section_header)

    editorial_groups = {g['clean_name']: g for g in structure['editorial_media_groups']}
    for name in EDITORIAL_MEDIA_ORDER:
        if name in editorial_groups:
            editorial_groups[name]['first_item'] = remove_reporter_phrases(editorial_groups[name]['first_item'])
            for item in editorial_groups[name]['additional_items']:
                item['text'] = remove_reporter_phrases(item['text'])
            add_media_group_to_document(new_doc, editorial_groups[name])


    all_content = []
    for content in structure['other_content']:
        if content['type'] == 'section_header' and content['section'] == 'editorial':
            continue
        all_content.append(('other', content))
    for section_name, articles in structure['sections'].items():
        for article in articles:
            all_content.append(('article', article))
    all_content.sort(key=lambda x: x[1].get('index', x[1].get('start_index', 0)))
    
    previous_was_content = False
    last_article_idx = -1
    for idx, (content_type, content_data) in enumerate(all_content):
        if content_type == 'other':
            if content_data['type'] == 'section_header':
                add_section_header_to_doc(new_doc, content_data['text'])
            else:
                clean_text = remove_reporter_phrases(content_data['text'])
                p = new_doc.add_paragraph(clean_text)
                format_content_paragraph(p)
            previous_was_content = True
        elif content_type == 'article':
            content_data['text'] = remove_reporter_phrases(content_data['text'])
            add_article_to_document(new_doc, content_data, previous_was_content)
            previous_was_content = True
            last_article_idx = idx

    if last_article_idx != -1:
        add_end_marker(new_doc)

    new_doc.save(output_path)
    return output_path

# =============================================================================
# WEB SCRAPING FUNCTIONS (Updated from wiser_scrape.py)
# =============================================================================

from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager

@retry_step
def setup_webdriver(**kwargs):
    headless = kwargs.get('headless')
    st_module = kwargs.get('st_module')
    
    try:
        if st_module:
            st_module.write("Setting up Chrome options...")
            
        options = webdriver.ChromeOptions()
        if headless:
            options.add_argument("--headless")
        
        # Stability options for Streamlit Cloud
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-gpu")
        options.add_argument("--remote-debugging-port=9222")
        
        # Let Selenium Manager handle driver management automatically
        if st_module:
            st_module.write("Using Selenium Manager for automatic driver management...")
            
        driver = webdriver.Chrome(options=options)  # No service parameter needed
        driver.get(WISERS_URL)
        if st_module:
            st_module.write("✅ WebDriver setup complete.")
        return driver
        
    except Exception as e:
        if st_module:
            st_module.error(f"WebDriver setup failed: {e}")
        return None

        



@retry_step
def perform_login(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    group_name = kwargs.get('group_name')
    username = kwargs.get('username')
    password = kwargs.get('password')
    api_key = kwargs.get('api_key')
    st_module = kwargs.get('st_module')
     

    # --- Logic to fill the form ---
    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'input[data-qa-ci="groupid"]'))).send_keys(group_name)
    driver.find_element(By.CSS_SELECTOR, 'input[data-qa-ci="userid"]').send_keys(username)
    driver.find_element(By.CSS_SELECTOR, 'input[data-qa-ci="password"]').send_keys(password)
    
    # --- Captcha solving logic ---
    try:
        captcha_img = driver.find_element(By.CSS_SELECTOR, 'img.CaptchaField__CaptchaImage-hffgxm-5')
        captcha_src = captcha_img.get_attribute('src')
        img_data = base64.b64decode(captcha_src.split(',')[1])
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_captcha:
            tmp_captcha.write(img_data)
            tmp_captcha_path = tmp_captcha.name
        
        solver = TwoCaptcha(api_key)
        captcha_text = solver.normal(tmp_captcha_path)['code']
        os.remove(tmp_captcha_path)
        
        driver.find_element(By.CSS_SELECTOR, 'input.CaptchaField__Input-hffgxm-4').send_keys(captcha_text)
    except Exception as captcha_error:
        # If captcha fails here, raise immediately to trigger retry
        raise Exception(f"Failed during 2Captcha solving process: {captcha_error}")

    login_btn = driver.find_element(By.CSS_SELECTOR, 'input[data-qa-ci="button-login"]')
    login_btn.click()

    # --- NEW: Intelligent Login Verification ---
    # Wait for either the dashboard (success) or an error message (failure) to appear
    try:
        WebDriverWait(driver, 10).until(
            EC.any_of(
                EC.element_to_be_clickable((By.CSS_SELECTOR, 'div.sc-1kg7aw5-0.dgeiTV > button')), # Success element
                EC.visibility_of_element_located((By.CSS_SELECTOR, 'div.NewContent__StyledNewErrorCode-q19ga1-5')) # Failure element
            )
        )
    except TimeoutException:
        raise Exception("Login verification failed: The page did not load the dashboard or a known error message.")

    # Check if an error message is present
    try:
        error_element = driver.find_element(By.CSS_SELECTOR, 'div.NewContent__StyledNewErrorCode-q19ga1-5')
        error_text = error_element.text.strip()
        
        # Determine the specific error and create a user-friendly message
        if "User over limit" in error_text:
            msg = "Login Failed: The account has reached its login limit. It is likely already logged in elsewhere."
        elif "captcha error" in error_text:
            msg = "Login Failed: The captcha code was incorrect."
        elif "Sorry, your login details are incorrect, please try again." in error_text:
            msg = "Login Failed: Incorrect Group, Username, or Password."
        else:
            msg = f"Login Failed: An unrecognized error appeared: '{error_text}'"

        # Display the specific error in Streamlit and raise an exception to trigger the retry decorator
        if st_module:
            st_module.warning(msg)
        raise Exception(msg)
        
    except NoSuchElementException:
        # If the error element wasn't found, the login must have been successful.
        if st_module:
            st_module.write("✅ Login successfully verified.")
        # No error, so the function completes and the script continues.
        return
        


@retry_step
def close_tutorial_modal_ROBUST(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    status_text = kwargs.get('status_text')
    st = kwargs.get('st_module')
    
    status_text.text("Attempting to close tutorial modal...")
    try:
        close_btn = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, '#app-userstarterguide-0 button.close')))
        ActionChains(driver).move_to_element(close_btn).click(close_btn).perform()
        time.sleep(2)
        wait.until(EC.invisibility_of_element_located((By.ID, 'app-userstarterguide-0')))
        status_text.text("Modal closed successfully!")
    except TimeoutException:
        status_text.text("Modal did not appear or was already closed.")
    except Exception as e:
        st.warning(f"Modal could not be closed. Continuing... Error: {e}")

@retry_step
def switch_language_to_traditional_chinese(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st = kwargs.get('st_module')
     

    #waffle_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'div.sc-1kg7aw5-0.dgeiTV > retryTest')))
    waffle_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'div.sc-1kg7aw5-0.dgeiTV > button')))
    waffle_button.click()
    lang_toggle = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'li.wo__header__nav__navbar__item.dropdown > a.dropdown-toggle')))
    driver.execute_script("arguments[0].click();", lang_toggle)
    trad_chinese_link = wait.until(EC.element_to_be_clickable((By.XPATH, '//a[span[text()="繁體中文"]]')))
    trad_chinese_link.click()
    wait.until(EC.staleness_of(waffle_button))
    time.sleep(3)
    return True

@retry_step
def perform_author_search(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    author_name = kwargs.get('author')
    st = kwargs.get('st_module')
     

    toggle_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'div.toggle-collapse[data-toggle="collapse"]')))
    driver.execute_script("arguments[0].click();", toggle_button)
    time.sleep(3)
    my_media_dropdown_toggle = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'button.btn-naked.dropdown-toggle[data-toggle="dropdown"]')))
    my_media_dropdown_toggle.click()
    time.sleep(3)
    hongkong_option = wait.until(EC.element_to_be_clickable((By.XPATH, '//label[span[text()="各大香港報章"]]')))
    hongkong_option.click()
    time.sleep(3)
    author_input = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 'input.form-control[placeholder="作者"]')))
    author_input.clear()
    author_input.send_keys(author_name)
    search_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'button#toggle-query-execute.btn.btn-primary')))
    search_button.click()

@retry_step
def wait_for_search_results(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st_module = kwargs.get('st_module')
     

    # Wait for a common container element to ensure the page has responded.
    # This is more reliable than a fixed sleep.
    try:
        wait.until(EC.presence_of_element_located((
            By.CSS_SELECTOR, 
            # This selector waits for EITHER a results list OR a known no-results message
            'div.list-group, div.list-group-item, .no-results, [class*="empty"]'
        )))
    except TimeoutException:
        # If nothing loads at all, let the decorator handle it.
        raise TimeoutException("Page did not load any known content after search.")

    time.sleep(1) # Brief, tactical pause for any final JS rendering.

    # --- 1. Check for any sign of success (results) ---
    result_selectors = [
        'div.list-group-item.no-excerpt', 
        'div.list-group-item', 
        '.article-main'
    ]
    for selector in result_selectors:
        if driver.find_elements(By.CSS_SELECTOR, selector):
            if st_module:
                st_module.write("✅ Search results found.")
            return True

    # --- 2. Check for any sign of a clean failure (no results) ---
    no_results_selectors = [
        ".no-results", 
        "[class*='no-result']", 
        "[class*='empty']"
    ]
    for selector in no_results_selectors:
        if driver.find_elements(By.CSS_SELECTOR, selector):
            if st_module:
                st_module.warning("ℹ️ No results found for this query.")
            return False

    # --- 3. Handle the ambiguous case ---
    # If we get here, the page loaded something, but it was neither a result 
    # nor a 'no results' message. This is an unexpected state.
    # Raise an exception to trigger a retry.
    raise Exception("Search page loaded, but content was unrecognized.")


@retry_step
def click_first_result(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    original_window = kwargs.get('original_window')
    st = kwargs.get('st_module')
     

    first_article_link = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'div.list-group .list-group-item h4 a')))
    first_article_link.click()
    wait.until(EC.number_of_windows_to_be(2))
    for window_handle in driver.window_handles:
        if window_handle != original_window:
            driver.switch_to.window(window_handle)
            break


@retry_step
def go_back_to_search_form(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st = kwargs.get('st_module')
     

    re_search_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'div.media-left > a[href="/wevo/home"]')))
    re_search_button.click()
    wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'button#toggle-query-execute.btn.btn-primary')))
    time.sleep(3)
    return True

@retry_step
def parse_media_info_for_author(**kwargs):
    subheading_text = kwargs.get('subheading_text')
    author_name = kwargs.get('author_name')
    st = kwargs.get('st_module')
     

    media_part = subheading_text.split('|')[0].strip()
    page_match = re.search(r'([A-Z]\d{2})', media_part)
    if page_match:
        page_number = page_match.group(1)
        media_name_part = media_part[:page_match.start()].strip()
        mapped_name = next((v for k, v in MEDIA_NAME_MAPPINGS.items() if k in media_name_part), media_name_part)
        return f"{mapped_name} {page_number} {author_name}："

@retry_step
def scrape_author_article_content(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    author_name = kwargs.get('author_name')
    st = kwargs.get('st_module')
     

    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div.article-detail')))
    time.sleep(3)
    title = driver.find_element(By.CSS_SELECTOR, 'h3').text.strip()
    subheading_text = driver.find_element(By.CSS_SELECTOR, 'div.article-subheading').text.strip()
    media_info = parse_media_info_for_author(subheading_text=subheading_text,author_name=author_name,st_module=st)
    paragraphs = [p.text.strip() for p in driver.find_elements(By.CSS_SELECTOR, 'div.description p') if p.text.strip()]
    if paragraphs:
        formatted_first_paragraph = f"{media_info}{paragraphs[0]}"
        full_content = [formatted_first_paragraph] + paragraphs[1:]
        formatted_content_body = '\n\n'.join(full_content)
        final_output = f"{title}\n\n{formatted_content_body}"
    else:
        final_output = title
    return {'title': title, 'content': final_output}


@retry_step
def run_newspaper_editorial_task(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st = kwargs.get('st_module')

    dropdown_toggle = wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "li.dropdown-usersavedquery > a.dropdown-toggle")))
    dropdown_toggle.click()
    time.sleep(3)
    edit_saved_search_btn = wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "button[data-target='#modal-saved-search-ws6']")))
    edit_saved_search_btn.click()
    wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#modal-saved-search-ws6")))
    time.sleep(3)
    editorial_item = wait.until(
        EC.element_to_be_clickable((By.XPATH, "//ul[@class='list-group']//h5[text()='社評']/ancestor::li")))
    editorial_item.click()
    time.sleep(3)

    search_btn = None
    selectors = [(By.CSS_SELECTOR, "div.modal-footer .btn-default:last-child"),
                 (By.XPATH, "//div[@class='modal-footer']//button[text()='搜索']")]
    for selector_type, selector in selectors:
        try:
            search_btn = wait.until(EC.element_to_be_clickable((selector_type, selector)))
            break
        except TimeoutException:
            continue
    if search_btn:
        search_btn.click()
    else:
        driver.execute_script("""
            var buttons = document.querySelectorAll('div.modal-footer button');
            for (var i = 0; i < buttons.length; i++) {
                if (buttons[i].textContent.trim() === '搜索') {
                    buttons[i].click(); break;
                }
            }""")
    wait.until(EC.invisibility_of_element_located((By.CSS_SELECTOR, "#modal-saved-search-ws6")))

    if wait_for_search_results(driver=driver, wait=wait, st_module=st):
        # NEW: Scroll to load all content, then also wait for AJAX to finish
        scroll_to_load_all_content(driver=driver, st_module=st)
        wait_for_ajax_complete(driver, timeout=10)

        # Now collect all results, with retries
        articles = []
        for retry in range(3):
            results = driver.find_elements(By.CSS_SELECTOR, 'div.list-group-item.no-excerpt')
            if st:
                st.write(f"[Editorial Scrape] Attempt {retry+1}: {len(results)} items found.")
            for result in results:
                try:
                    title = result.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a').text.strip()
                    media_name_raw = result.find_element(By.CSS_SELECTOR, 'small a').text.strip()
                    mapped_name = next((v for k, v in MEDIA_NAME_MAPPINGS.items() if k in media_name_raw), media_name_raw)
                    article = {'media': mapped_name, 'title': title}
                    if article not in articles:
                        articles.append(article)
                except Exception:
                    continue
            if len(articles) > 0:
                break
            time.sleep(2)
        return articles

    return []



@retry_step
def scroll_to_load_all_content(**kwargs):
    """Scroll to bottom to trigger lazy loading of all editorial content."""
    driver = kwargs.get('driver')
    st_module = kwargs.get('st_module')

    max_attempts = 10  # Avoid infinite loop on buggy sites
    last_height = driver.execute_script("return document.body.scrollHeight")

    for attempt in range(max_attempts):
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)
        new_height = driver.execute_script("return document.body.scrollHeight")
        if st_module:
            st_module.write(f"[Scroll] Pass {attempt+1}: Height {new_height}")
        if new_height == last_height:
            break
        last_height = new_height
    if st_module:
        st_module.write("Scrolling finished (all editorial content should be loaded now).")
    return True

def wait_for_ajax_complete(driver, timeout=10):
    """Wait for jQuery AJAX calls to complete if jQuery is present."""
    try:
        WebDriverWait(driver, timeout).until(
            lambda d: d.execute_script("return jQuery.active == 0") if d.execute_script("return typeof jQuery != 'undefined'") else True
        )
    except Exception:
        pass

@retry_step
def run_scmp_editorial_task(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st = kwargs.get('st_module')

    toggle_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'div.toggle-collapse[data-toggle="collapse"]')))
    driver.execute_script("arguments[0].click();", toggle_button)
    time.sleep(2)
    my_media_dropdown_toggle = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'button.btn-naked.dropdown-toggle[data-toggle="dropdown"]')))
    my_media_dropdown_toggle.click()
    time.sleep(1)
    hongkong_option = wait.until(EC.element_to_be_clickable((By.XPATH, '//label[span[text()="各大香港報章"]]')))
    hongkong_option.click()
    time.sleep(1)
    author_input = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 'input.form-control[placeholder="欄目"]')))
    author_input.clear()
    author_input.send_keys("editorial")
    search_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'button#toggle-query-execute.btn.btn-primary')))
    search_button.click()

    # Scroll and wait for AJAX after search to maximize completeness
    scroll_to_load_all_content(driver=driver, st_module=st)
    wait_for_ajax_complete(driver, timeout=10)

    if wait_for_search_results(driver=driver, wait=wait, st_module=st):
        articles = []
        for retry in range(3):
            results = driver.find_elements(By.CSS_SELECTOR, 'div.list-group-item.no-excerpt')
            if st:
                st.write(f"[SCMP Editorial Scrape] Attempt {retry+1}: {len(results)} items found.")
            for result in results:
                try:
                    title = result.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a').text.strip()
                    media_name_raw = result.find_element(By.CSS_SELECTOR, 'small a').text.strip()
                    mapped_name = next((v for k, v in MEDIA_NAME_MAPPINGS.items() if k in media_name_raw), None)
                    if mapped_name == 'SCMP':
                        article = {'media': 'SCMP', 'title': title}
                        if article not in articles:
                            articles.append(article)
                except Exception:
                    continue
            if len(articles) > 0:
                break
            time.sleep(2)
        return articles
    return []


@retry_step
def create_docx_report(**kwargs):
    author_articles_data = kwargs.get('author_articles_data')
    editorial_data = kwargs.get('editorial_data')
    author_list = kwargs.get('author_list')
    output_path = kwargs.get('output_path')
    st = kwargs.get('st_module')
     
    from docx import Document

    doc = Document()
    doc.add_heading('指定作者社評', level=1)
    doc.add_paragraph()
    for author in author_list:
        article = author_articles_data.get(author)
        title = article['title'] if article else ""
        doc.add_paragraph(f"{author}：{title}")
    doc.add_paragraph()
    is_first_article = True
    for author in author_list:
        article = author_articles_data.get(author)
        if article and article.get('content'):
            if not is_first_article:
                doc.add_paragraph()
            for paragraph_text in article['content'].split('\n\n'):
                doc.add_paragraph(paragraph_text)
            is_first_article = False

    if editorial_data:
        doc.add_page_break()
        doc.add_heading('報章社評', level=1)
        doc.add_paragraph()
        grouped_editorials = defaultdict(list)
        for article in editorial_data:
            grouped_editorials[article['media']].append(article['title'])

        for media, titles in grouped_editorials.items():
            if len(titles) == 1:
                doc.add_paragraph(f"{media}：{titles[0]}")
            else:
                doc.add_paragraph(f"{media}：1. {titles[0]}")
                for i, title in enumerate(titles[1:], start=2):
                    p = doc.add_paragraph()
                    p.add_run(f"\t{i}. {title}")
    
    doc.save(output_path)
    return output_path

@retry_step
def scroll_to_load_all_content(**kwargs):
    """Scroll to bottom to trigger lazy loading of all editorial content."""
    driver = kwargs.get('driver')
    st_module = kwargs.get('st_module')

    max_attempts = 10  # Avoid infinite loop on buggy sites

    last_height = driver.execute_script("return document.body.scrollHeight")

    for attempt in range(max_attempts):
        # Scroll down to bottom
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)  # Wait for new content to load
        new_height = driver.execute_script("return document.body.scrollHeight")

        if st_module:
            st_module.write(f"[Scroll] Pass {attempt+1}: Height {new_height}")

        if new_height == last_height:
            break  # No more content to load

        last_height = new_height

    if st_module:
        st_module.write("Scrolling finished (all editorial content should be loaded now).")
    return True

def wait_for_ajax_complete(driver, timeout=10):
    """Wait for jQuery AJAX calls to complete if jQuery is present."""
    try:
        WebDriverWait(driver, timeout).until(
            lambda d: d.execute_script("return jQuery.active == 0") if d.execute_script("return typeof jQuery != 'undefined'") else True
        )
    except Exception:
        pass  # If jQuery not defined, or page doesn't use it, just continue.


@retry_step
def logout(**kwargs):
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st = kwargs.get('st_module')
     

    waffle_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'div.sc-1kg7aw5-0.dgeiTV > button')))
    waffle_button.click()
    time.sleep(1)
    logout_link = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "li.wo__header__nav__navbar__item:not(.dropdown) a")))
    logout_link.click()
    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'input[data-qa-ci="groupid"]')))

def robust_logout_request(driver, st_module=None):
    """Send the robust logout API GET request to forcibly close session."""
    import requests
    from selenium.webdriver.chrome.webdriver import WebDriver
    import time

    if not driver:
        if st_module:
            st_module.warning("robust_logout_request: driver is None")
        return

    if not isinstance(driver, WebDriver):
        if st_module:
            st_module.warning("robust_logout_request requires a selenium WebDriver instance.")
        return

    try:
        # Extract session cookies from Selenium driver
        selenium_cookies = driver.get_cookies()
        if st_module:
            st_module.write(f"Found {len(selenium_cookies)} cookies from driver")
            
        session_cookies = {cookie['name']: cookie['value'] for cookie in selenium_cookies}

        # Get current timestamp for the logout URL
        current_timestamp = int(time.time() * 1000)
        
        robust_logout_url = (
            "https://wisesearch6.wisers.net/wevo/api/AccountService;criteria=%7B%22groupId%22%3A%22SPRG1%22%2C"
            "%22userId%22%3A%22AsiaNet1%22%2C%22deviceType%22%3A%22web%22%2C%22deviceId%22%3A%22%22%7D;"
            f"path=logout;timestamp={current_timestamp};updateSession=true"
            "?returnMeta=true"
        )

        headers = {
            "accept": "*/*",
            "accept-language": "en-US,en;q=0.9",
            "sec-ch-ua": '"Not)A;Brand";v="8", "Chromium";v="138", "Google Chrome";v="138"',
            "sec-ch-ua-mobile": "?0",
            "sec-ch-ua-platform": '"Windows"',
            "sec-fetch-dest": "empty",
            "sec-fetch-mode": "cors",
            "sec-fetch-site": "same-origin",
            "x-requested-with": "XMLHttpRequest"
        }

        if st_module:
            st_module.write("Sending robust logout request...")

        response = requests.get(robust_logout_url, headers=headers, cookies=session_cookies, timeout=10)
        
        if st_module:
            st_module.write(f"Logout response status: {response.status_code}")
            st_module.write(f"Logout response text: {response.text[:200]}...")

        if response.ok:
            if st_module:
                st_module.write("✅ Robust logout request sent successfully.")
        else:
            if st_module:
                st_module.warning(f"Robust logout request failed with status: {response.status_code}")
                
    except Exception as e:
        if st_module:
            st_module.warning(f"Exception during robust logout request: {e}")
            import traceback
            st_module.code(traceback.format_exc())

# =============================================================================
# WEB SCRAPING FUNCTIONS For International News
# =============================================================================

@retry_step
def run_international_news_task(**kwargs):
    """Search for international news articles"""
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    st = kwargs.get('st_module')
    
    # Click on saved search dropdown
    dropdown_toggle = wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "li.dropdown-usersavedquery > a.dropdown-toggle")))
    dropdown_toggle.click()
    time.sleep(3)

    # Open saved search modal
    edit_saved_search_btn = wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "button[data-target='#modal-saved-search-ws6']")))
    edit_saved_search_btn.click()
    wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "#modal-saved-search-ws6")))
    time.sleep(3)

    # Look for "國際新聞" in the saved searches
    international_item = wait.until(
        EC.element_to_be_clickable((By.XPATH, "//ul[@class='list-group']//h5[text()='國際新聞']/ancestor::li")))
    international_item.click()
    time.sleep(3)

    # Click search button
    search_btn = None
    selectors = [(By.CSS_SELECTOR, "div.modal-footer .btn-default:last-child"),
                (By.XPATH, "//div[@class='modal-footer']//button[text()='搜索']")]
    
    for selector_type, selector in selectors:
        try:
            search_btn = wait.until(EC.element_to_be_clickable((selector_type, selector)))
            break
        except TimeoutException:
            continue

    if search_btn:
        search_btn.click()
    else:
        driver.execute_script("""
            var buttons = document.querySelectorAll('div.modal-footer button');
            for (var i = 0; i < buttons.length; i++) {
                if (buttons[i].textContent.trim() === '搜索') {
                    buttons[i].click(); break;
                }
            }""")

    # Wait for modal to close
    wait.until(EC.invisibility_of_element_located((By.CSS_SELECTOR, "#modal-saved-search-ws6")))
    
    if wait_for_search_results(driver=driver, wait=wait, st_module=st):
        # Scroll to load all content and wait for AJAX
        scroll_to_load_all_content(driver=driver, st_module=st)
        wait_for_ajax_complete(driver, timeout=10)
        
        # Get all article links for scraping
        articles_data = []
        
        for retry in range(3):
            results = driver.find_elements(By.CSS_SELECTOR, 'div.list-group-item.no-excerpt')
            if st:
                st.write(f"[International News Scrape] Attempt {retry+1}: {len(results)} items found.")
            
            # Limit to around 80-100 articles as requested
            results = results[:100]
            
            for i, result in enumerate(results):
                try:
                    title_element = result.find_element(By.CSS_SELECTOR, 'h4.list-group-item-heading a')
                    title = title_element.text.strip()
                    article_url = title_element.get_attribute('href')
                    
                    # Get media name
                    try:
                        media_name_raw = result.find_element(By.CSS_SELECTOR, 'small a').text.strip()
                        mapped_name = next((v for k, v in MEDIA_NAME_MAPPINGS.items() if k in media_name_raw), media_name_raw)
                    except:
                        mapped_name = "Unknown"
                    
                    articles_data.append({
                        'title': title,
                        'url': article_url,
                        'media': mapped_name,
                        'index': i
                    })
                    
                except Exception as e:
                    if st:
                        st.warning(f"Error extracting article {i}: {e}")
                    continue
            
            if len(articles_data) > 0:
                break
            time.sleep(2)
        
        return articles_data
    
    return []

@retry_step
def scrape_international_article_detail(**kwargs):
    """Scrape individual international news article content"""
    driver = kwargs.get('driver')
    wait = kwargs.get('wait')
    article_url = kwargs.get('article_url')
    original_window = kwargs.get('original_window')
    st = kwargs.get('st_module')
    
    # Open article in new tab
    driver.execute_script(f"window.open('{article_url}', '_blank');")
    
    # Wait for new window and switch to it
    wait.until(EC.number_of_windows_to_be(2))
    for window_handle in driver.window_handles:
        if window_handle != original_window:
            driver.switch_to.window(window_handle)
            break
    
    # Wait for article content to load
    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div.article-detail')))
    time.sleep(2)
    
    try:
        # Get title
        title = driver.find_element(By.CSS_SELECTOR, 'h3').text.strip()
        
        # Get subheading/metadata
        try:
            subheading_text = driver.find_element(By.CSS_SELECTOR, 'div.article-subheading').text.strip()
        except:
            subheading_text = ""
        
        # Get article content paragraphs
        paragraphs = []
        content_elements = driver.find_elements(By.CSS_SELECTOR, 'div.description p')
        for p in content_elements:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        
        # Combine content
        content_body = '\n\n'.join(paragraphs) if paragraphs else ""
        
        article_data = {
            'title': title,
            'subheading': subheading_text,
            'content': content_body,
            'full_text': f"{title}\n\n{subheading_text}\n\n{content_body}" if content_body else title
        }
        
        return article_data
        
    except Exception as e:
        if st:
            st.warning(f"Error scraping article content: {e}")
        return None
    
    finally:
        # Close current tab and return to original window
        driver.close()
        driver.switch_to.window(original_window)

@retry_step
def create_international_news_report(**kwargs):
    """Create Word document report for international news"""
    articles_data = kwargs.get('articles_data')
    output_path = kwargs.get('output_path')
    st = kwargs.get('st_module')
    
    from docx import Document
    
    doc = Document()
    setup_document_fonts(doc)
    
    # Add title
    doc.add_heading('國際新聞摘要', level=1)
    doc.add_paragraph()
    
    # Add date
    today_str = datetime.now().strftime("%Y年%m月%d日")
    date_para = doc.add_paragraph(f"日期：{today_str}")
    date_para.add_run().add_break()
    
    # Add articles
    for i, article in enumerate(articles_data, 1):
        if article and article.get('full_text'):
            # Add article number and title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{i}. {article['title']}")
            title_run.bold = True
            
            # Add media source if available
            if article.get('media'):
                media_para = doc.add_paragraph(f"來源：{article['media']}")
                media_para.style = doc.styles['Normal']
            
            # Add content
            if article.get('content'):
                for paragraph_text in article['content'].split('\n\n'):
                    if paragraph_text.strip():
                        doc.add_paragraph(paragraph_text.strip())
            
            # Add spacing between articles
            doc.add_paragraph()
    
    # Add end marker
    add_end_marker(doc)
    
    doc.save(output_path)
    return output_path

# =============================================================================
# STREAMLIT APP
# =============================================================================

def main():
    st.set_page_config(page_title="Document Processing Tool", layout="wide")
    
    try:
        if not st.secrets.get("wisers", {}).get("api_key"):
            st.warning("⚠️ Secrets not configured. Manual input will be required for web scraping.")
    except Exception as e:
        if isinstance(e, st.errors.StreamlitAPIException):
            st.warning("⚠️ Secrets not configured locally. Manual input required.")
        else:
            st.warning(f"Error checking secrets: {e}")

    st.title("AsiaNet Document Processing Tool")
    st.markdown("Choose between document formatting or web scraping functionality")

    # Update tabs to include the new International News tab
    tab1, tab2, tab3 = st.tabs(["📄 Document Formatting", "🌐 Web Scraping & Reporting", "🌍 International News"])

    with tab1:
        # Existing document formatting code remains the same
        st.header("Document Formatting")
        st.markdown("Upload your Word document to get it formatted automatically")

        uploaded_file = st.file_uploader("Choose a Word document", type=['docx'])

        if uploaded_file is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_file_path = tmp_file.name

            try:
                st.write("Processing your document...")
                progress = st.progress(0)
                
                progress.progress(25, text="Extracting document structure...")
                structure = extract_document_structure(tmp_file_path)
                
                progress.progress(50, text="Rebuilding document from structure...")
                formatted_file = rebuild_document_from_structure(tmp_file_path)
                
                progress.progress(75, text="Applying headers and footers...")
                logo_path = "AsiaNet_logo.png"
                if os.path.exists(logo_path):
                    doc = Document(formatted_file)
                    add_first_page_header(doc, logo_path)
                    add_first_page_footer(doc)
                    add_subsequent_pages_header(doc)
                    add_subsequent_pages_footer(doc)
                    doc.save(formatted_file)
                
                progress.progress(100, text="Formatting complete!")

                with open(formatted_file, 'rb') as f:
                    st.download_button(
                        label="📥 Download Formatted Document",
                        data=f.read(),
                        file_name=f"formatted_{uploaded_file.name}",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                st.success("Document processed successfully!")

            except Exception as e:
                st.error(f"Error processing document: {str(e)}")
            finally:
                if 'tmp_file_path' in locals() and os.path.exists(tmp_file_path):
                    os.remove(tmp_file_path)
                if 'formatted_file' in locals() and os.path.exists(formatted_file):
                    os.remove(formatted_file)

    with tab2:
        # Existing web scraping code remains exactly the same
        st.header("Web Scraping and Report Generation")
        st.markdown("Scrape articles by specified authors and newspaper editorials, then generate a combined Word report.")

        with st.expander("⚙️ Scraping Configuration", expanded=True):
            col1, col2 = st.columns(2)

            with col1:
                try:
                    group_name = st.secrets["wisers"]["group_name"]
                    username = st.secrets["wisers"]["username"]
                    password = st.secrets["wisers"]["password"]
                    st.success("✅ Credentials loaded from secrets")
                    st.info(f"Group: {group_name}\n\nUsername: {username}\n\nPassword: ****")
                except (KeyError, AttributeError, st.errors.StreamlitAPIException):
                    st.warning("⚠️ Secrets not found. Please enter credentials manually:")
                    group_name = st.text_input("Group Name", value="SPRG1")
                    username = st.text_input("Username", placeholder="Enter username")
                    password = st.text_input("Password", type="password", placeholder="Enter password")

            with col2:
                try:
                    api_key = st.secrets["wisers"]["api_key"]
                    st.success(f"✅ 2Captcha API Key loaded: {api_key[:8]}...")
                except (KeyError, AttributeError, st.errors.StreamlitAPIException):
                    st.warning("⚠️ API key not found in secrets")
                    api_key = st.text_input("2Captcha API Key", type="password", placeholder="Enter API key")

        authors_input = st.text_area("Authors to Search (one per line)",
                                   value="李先知\n余錦賢\n傅流螢",
                                   help="Enter one author name per line. The script will search for the latest article from each.")

        st.sidebar.header("Debugging Options")
        st.sidebar.markdown("---")
        run_headless = st.checkbox("Run in headless mode (faster, no visible browser)", value=True)
        keep_browser_open = st.sidebar.checkbox("Keep browser open after script finishes/fails")

        if st.button("🚀 Start Scraping and Generate Report", type="primary"):
            # All existing scraping logic remains the same...
            if not all([group_name, username, password, api_key]):
                st.error("❌ Please provide all required credentials and the API key to proceed.")
                st.stop()

            authors_list = [author.strip() for author in authors_input.split('\n') if author.strip()]
            if not authors_list:
                st.error("❌ Please enter at least one author to search.")
                st.stop()

            progress_bar = st.progress(0)
            status_text = st.empty()
            driver = None

            try:
                status_text.text("Setting up web driver...")
                driver = setup_webdriver(headless=run_headless, st_module=st)

                if driver is None:
                    st.error("Driver setup failed, cannot continue. See logs above for details.")
                    st.stop()

                wait = WebDriverWait(driver, 20)
                progress_bar.progress(5, text="Driver ready. Logging in...")

                perform_login(driver=driver,wait=wait,group_name=group_name,username=username,password=password,api_key=api_key,st_module=st)

                progress_bar.progress(10, text="Login successful. Finalizing setup...")
                time.sleep(5)

                close_tutorial_modal_ROBUST(driver=driver, wait=wait, status_text=status_text, st_module=st)
                switch_language_to_traditional_chinese(driver=driver, wait=wait, st_module=st)

                progress_bar.progress(15, text="Language set. Starting author search...")

                original_window = driver.current_window_handle
                author_articles_data = {}
                total_steps = len(authors_list) + 3
                progress_increment = 70 / total_steps

                # Author search loop - existing code
                for i, author in enumerate(authors_list):
                    current_progress = 15 + (i * progress_increment)
                    status_text.text(f"({i+1}/{len(authors_list)}) Searching for author: {author}...")
                    progress_bar.progress(int(current_progress), text=f"Searching for {author}")

                    perform_author_search(driver=driver, wait=wait, author=author, st_module=st)

                    if wait_for_search_results(driver=driver, wait=wait, st_module=st):
                        click_first_result(driver=driver, wait=wait, original_window=original_window, st_module=st)
                        scraped_data = scrape_author_article_content(driver=driver, wait=wait, author_name=author, st_module=st)
                        author_articles_data[author] = scraped_data

                        st.write("Closing article tab and returning to search results...")
                        driver.close()
                        driver.switch_to.window(original_window)
                    else:
                        author_articles_data[author] = None
                        st.info(f"No results found for {author}.")

                    go_back_to_search_form(driver=driver, wait=wait, st_module=st)

                # Editorial tasks - existing code
                final_author_progress = 15 + (len(authors_list) * progress_increment)
                progress_bar.progress(int(final_author_progress), text="Scraping newspaper editorials...")
                status_text.text("Scraping newspaper editorials (from saved search)...")

                editorial_data = run_newspaper_editorial_task(driver=driver, wait=wait, st_module=st)
                if editorial_data is None: editorial_data = []

                st.write("Returning to main search form for SCMP task...")
                go_back_to_search_form(driver=driver, wait=wait, st_module=st)

                progress_bar.progress(int(final_author_progress + progress_increment), text="Scraping SCMP editorials...")
                status_text.text("Scraping SCMP editorials (manual search)...")

                scmp_editorial_data = run_scmp_editorial_task(driver=driver, wait=wait, st_module=st)
                if scmp_editorial_data:
                    editorial_data.extend(scmp_editorial_data)

                # Report generation - existing code
                progress_bar.progress(int(final_author_progress + 2 * progress_increment), text="Generating Word document...")
                status_text.text("Creating final Word report...")

                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_report:
                    output_path = create_docx_report(author_articles_data=author_articles_data,editorial_data=editorial_data,
                                                   author_list=authors_list,output_path=tmp_report.name,st_module=st)

                progress_bar.progress(95, text="Report generated. Logging out...")
                status_text.text("Logging out...")

                logout(driver=driver, wait=wait, st_module=st)
                robust_logout_request(driver, st_module=st)

                with open(output_path, 'rb') as f:
                    st.download_button(
                        label="📥 Download Combined Report",
                        data=f.read(),
                        file_name=f"香港社評報告_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                progress_bar.progress(100, text="✅ Process complete!")
                status_text.success("✅ Scraping and report generation completed successfully!")

                st.subheader("📊 Scraped Content Summary")
                for author, data in author_articles_data.items():
                    st.write(f"**{author}**: {'Article found' if data else 'No article found'}")
                st.write(f"**Editorials**: Found {len(editorial_data)} total editorial articles.")

                st.success("✅ Scraping process completed successfully!")

            except Exception as e:
                st.error(f"❌ A critical error stopped the script: {str(e)}")
                st.code(traceback.format_exc())

            finally:
                try:
                    if 'driver' in locals() and driver:
                        if not keep_browser_open:
                            robust_logout_request(driver, st_module=st)
                        else:
                            st.warning("🤖 As requested, the browser window has been left open for inspection.")
                except Exception as cleanup_err:
                    st.error(f"Error in cleanup: {cleanup_err}")

    # NEW TAB 3: International News
    with tab3:
        st.header("International News Scraping")
        st.markdown("Scrape 80-100 pieces of international news articles and generate a Word report.")

        with st.expander("⚙️ International News Configuration", expanded=True):
            col1, col2 = st.columns(2)

            with col1:
                try:
                    group_name_intl = st.secrets["wisers"]["group_name"]
                    username_intl = st.secrets["wisers"]["username"]
                    password_intl = st.secrets["wisers"]["password"]
                    st.success("✅ Credentials loaded from secrets")
                    st.info(f"Group: {group_name_intl}\n\nUsername: {username_intl}\n\nPassword: ****")
                except (KeyError, AttributeError, st.errors.StreamlitAPIException):
                    st.warning("⚠️ Secrets not found. Please enter credentials manually:")
                    group_name_intl = st.text_input("Group Name", value="SPRG1", key="intl_group")
                    username_intl = st.text_input("Username", placeholder="Enter username", key="intl_username")
                    password_intl = st.text_input("Password", type="password", placeholder="Enter password", key="intl_password")

            with col2:
                try:
                    api_key_intl = st.secrets["wisers"]["api_key"]
                    st.success(f"✅ 2Captcha API Key loaded: {api_key_intl[:8]}...")
                except (KeyError, AttributeError, st.errors.StreamlitAPIException):
                    st.warning("⚠️ API key not found in secrets")
                    api_key_intl = st.text_input("2Captcha API Key", type="password", placeholder="Enter API key", key="intl_api")

        # International news specific settings
        max_articles = st.slider("Maximum articles to scrape", min_value=50, max_value=150, value=100, 
                                help="Limit the number of international news articles to scrape")

        st.sidebar.header("International News Options")
        st.sidebar.markdown("---")
        run_headless_intl = st.sidebar.checkbox("Run in headless mode (faster, no visible browser)", value=True, key="intl_headless")
        keep_browser_open_intl = st.sidebar.checkbox("Keep browser open after script finishes/fails", key="intl_keep_open")

        if st.button("🌍 Start International News Scraping", type="primary"):
            if not all([group_name_intl, username_intl, password_intl, api_key_intl]):
                st.error("❌ Please provide all required credentials and the API key to proceed.")
                st.stop()

            progress_bar = st.progress(0)
            status_text = st.empty()
            driver = None

            try:
                status_text.text("Setting up web driver for international news...")
                driver = setup_webdriver(headless=run_headless_intl, st_module=st)

                if driver is None:
                    st.error("Driver setup failed, cannot continue. See logs above for details.")
                    st.stop()

                wait = WebDriverWait(driver, 20)
                progress_bar.progress(5, text="Driver ready. Logging in...")

                perform_login(driver=driver, wait=wait, group_name=group_name_intl, 
                             username=username_intl, password=password_intl, 
                             api_key=api_key_intl, st_module=st)

                progress_bar.progress(10, text="Login successful. Finalizing setup...")
                time.sleep(5)

                close_tutorial_modal_ROBUST(driver=driver, wait=wait, status_text=status_text, st_module=st)
                switch_language_to_traditional_chinese(driver=driver, wait=wait, st_module=st)

                progress_bar.progress(20, text="Language set. Searching for international news...")

                # Get international news article list
                status_text.text("Searching for international news articles...")
                articles_list = run_international_news_task(driver=driver, wait=wait, st_module=st)

                if not articles_list:
                    st.warning("No international news articles found.")
                    st.stop()

                # Limit articles to max_articles
                articles_list = articles_list[:max_articles]
                st.info(f"Found {len(articles_list)} articles to scrape.")

                progress_bar.progress(30, text=f"Found {len(articles_list)} articles. Starting detailed scraping...")

                # Scrape individual articles
                scraped_articles = []
                original_window = driver.current_window_handle
                
                for i, article_info in enumerate(articles_list):
                    current_progress = 30 + (i * 50 / len(articles_list))
                    status_text.text(f"Scraping article {i+1}/{len(articles_list)}: {article_info['title'][:50]}...")
                    progress_bar.progress(int(current_progress), text=f"Scraping article {i+1}/{len(articles_list)}")

                    try:
                        article_content = scrape_international_article_detail(
                            driver=driver, wait=wait, article_url=article_info['url'],
                            original_window=original_window, st_module=st
                        )
                        
                        if article_content:
                            article_content['media'] = article_info['media']
                            scraped_articles.append(article_content)
                        
                    except Exception as e:
                        st.warning(f"Failed to scrape article {i+1}: {e}")
                        continue

                progress_bar.progress(80, text="Creating Word document report...")
                status_text.text("Generating international news report...")

                # Generate report
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_report:
                    output_path = create_international_news_report(
                        articles_data=scraped_articles, 
                        output_path=tmp_report.name, 
                        st_module=st
                    )

                progress_bar.progress(90, text="Report generated. Logging out...")
                status_text.text("Logging out...")

                logout(driver=driver, wait=wait, st_module=st)
                robust_logout_request(driver, st_module=st)

                # Provide download
                with open(output_path, 'rb') as f:
                    st.download_button(
                        label="📥 Download International News Report",
                        data=f.read(),
                        file_name=f"國際新聞報告_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                progress_bar.progress(100, text="✅ International news scraping complete!")
                status_text.success("✅ International news scraping completed successfully!")

                st.subheader("📊 International News Summary")
                st.write(f"**Total articles scraped**: {len(scraped_articles)}")
                
                # Show media breakdown
                media_count = {}
                for article in scraped_articles:
                    media = article.get('media', 'Unknown')
                    media_count[media] = media_count.get(media, 0) + 1
                
                for media, count in sorted(media_count.items()):
                    st.write(f"**{media}**: {count} articles")

                st.success("✅ International news scraping process completed successfully!")

            except Exception as e:
                st.error(f"❌ A critical error stopped the international news script: {str(e)}")
                st.code(traceback.format_exc())

            finally:
                try:
                    if 'driver' in locals() and driver:
                        if not keep_browser_open_intl:
                            robust_logout_request(driver, st_module=st)
                        else:
                            st.warning("🤖 As requested, the browser window has been left open for inspection.")
                except Exception as cleanup_err:
                    st.error(f"Error in cleanup: {cleanup_err}")

if __name__ == "__main__":
    main()